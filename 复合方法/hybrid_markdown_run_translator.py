import json
import os
import re
import threading
import time
import traceback
from copy import copy
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from tkinter import BOTH, END, LEFT, RIGHT, X, BooleanVar, Button, Canvas, Checkbutton, Entry, Frame, Label, OptionMenu, Scrollbar, StringVar, Text, Tk, filedialog, messagebox
from tkinter.ttk import Progressbar

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD

    DND_AVAILABLE = True
except Exception:
    TkinterDnD = None
    DND_FILES = None
    DND_AVAILABLE = False


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "输出"
CONFIG_PATH = APP_DIR / ".hybrid_config.json"

PROVIDER_DEFAULTS = {
    "OpenAI": {"base_url": "", "model": "gpt-4.1", "key_label": "OpenAI API Key"},
    "DeepSeek": {"base_url": "https://api.deepseek.com", "model": "deepseek-v4-flash", "key_label": "DeepSeek API Key"},
}
APP_VERSION = "v1.16"
ENGLISH_FONT_OPTIONS = ("Times New Roman", "Calibri")
CHINESE_FONT_OPTIONS = ("楷体_GB2312", "宋体")
DEFAULT_ENGLISH_FONT = "Times New Roman"
DEFAULT_CHINESE_FONT = "宋体"
DIGIT_FONT = "Times New Roman"
PROGRESS_TOTAL = 100
ENABLE_COMPANY_NAME_GLOSSARY = False
STRICT_NO_CJK_OUTPUT = True

SYSTEM_PROMPT = """You are a senior bilingual legal translator and document-format alignment specialist.
Translate Chinese legal contracts into precise formal legal English.
Preserve legal meaning, defined terms, numbering, dates, parties, amounts, placeholders, and clause references.
The final English output must not contain Chinese/CJK characters; translate or romanize proper names instead of using English (Chinese) format.
Return only valid JSON when JSON is requested."""

FINAL_AUDIT_SYSTEM_PROMPT = """You are a senior legal translation QA reviewer.
Your task is to quickly audit completed English translations of Chinese capital-markets legal documents under a strict English-only final-output policy.
Flag any block whose translation still contains Chinese/CJK characters, including legal prose, clause headings, table labels, TOC entries, company names, fund names, investor names, person names, addresses, trademarks, parentheticals, or mixed Chinese-English leftovers.
Do not allow English Name (Chinese Name) formatting. Translate or romanize Chinese proper names into English-only text.
Preserve source placeholder symbols such as 【】 and blanks, but translate any Chinese words inside placeholders.
Return only valid JSON."""

CAPITAL_MARKETS_LEGAL_RAG = """Capital markets legal English retrieval notes:
- This is a capital markets / private equity style legal contract. Prefer formal transactional drafting, not conversational English.
- Translate \u9644\u5f55 as Appendix and keep it consistent. Do not translate \u9644\u5f55 as Schedule.
- Translate \u9644\u4ef6 as Annex. Translate \u9644\u8868 as Schedule.
- Use Roman numerals for Chinese appendix numerals: \u9644\u5f55\u4e00 -> Appendix I; \u9644\u5f55\u4e8c -> Appendix II; \u9644\u5f55\u4e09\uff08A\uff09 -> Appendix III(A); \u9644\u5f55\u4e03 -> Appendix VII.
- Translate \u62ab\u9732\u51fd / \u62ab\u9732\u6e05\u5355 as Disclosure Schedule. If it is titled \u9644\u5f55\u4e94\uff08\u62ab\u9732\u51fd\uff09, use Appendix V (Disclosure Schedule).
- Translate \u4ea4\u5272 as Closing, \u4ea4\u5272\u65e5 as Closing Date, \u6700\u8fdf\u5b8c\u6210\u65e5 as Longstop Date, \u7b7e\u7f72\u65e5 as Signing Date.
- Translate \u672c\u6b21\u4ea4\u6613 as this Transaction, \u6295\u8d44\u6b3e as Investment Amount, \u76ee\u6807\u516c\u53f8 as Target Company, \u96c6\u56e2\u516c\u53f8 as Group Company / Group Companies.
- For RMB amounts expressed as \u4eba\u6c11\u5e01X\u4e07\u5143, convert to yuan in English: \u4eba\u6c11\u5e0112.7764\u4e07\u5143 -> RMB 127,764. Do not output Ten Thousand Yuan or million for \u4e07\u5143.
- Preserve source \u3010...\u3011 bracket placeholder symbols exactly. \u3010\u3011 stays \u3010\u3011, not [] or underscores; if a placeholder contains Chinese text, translate the text inside the brackets while keeping the \u3010...\u3011 form.
- Preserve defined-term capitalization once established. If the same Chinese term recurs, reuse the same English term.
- Use English punctuation in English output: straight quotes, half-width commas, periods, semicolons, colons, and parentheses, while preserving source \u3010...\u3011 placeholders."""

LEGAL_RAG_TERMS = [
    ("\u9644\u5f55", "\u9644\u5f55 = Appendix; use Roman numerals, e.g. Appendix I, Appendix II, Appendix III(A)."),
    ("\u9644\u4ef6", "\u9644\u4ef6 = Annex."),
    ("\u9644\u8868", "\u9644\u8868 = Schedule."),
    ("\u62ab\u9732\u51fd", "\u62ab\u9732\u51fd = Disclosure Schedule."),
    ("\u4ea4\u5272\u65e5", "\u4ea4\u5272\u65e5 = Closing Date."),
    ("\u4ea4\u5272", "\u4ea4\u5272 = Closing."),
    ("\u6700\u8fdf\u5b8c\u6210\u65e5", "\u6700\u8fdf\u5b8c\u6210\u65e5 = Longstop Date."),
    ("\u7b7e\u7f72\u65e5", "\u7b7e\u7f72\u65e5 = Signing Date."),
    ("\u672c\u6b21\u4ea4\u6613", "\u672c\u6b21\u4ea4\u6613 = this Transaction."),
    ("\u6295\u8d44\u6b3e", "\u6295\u8d44\u6b3e = Investment Amount."),
    ("\u76ee\u6807\u516c\u53f8", "\u76ee\u6807\u516c\u53f8 = Target Company."),
    ("\u96c6\u56e2\u516c\u53f8", "\u96c6\u56e2\u516c\u53f8 = Group Company / Group Companies."),
]

ORG_NAME_SUFFIXES = (
    "\u80a1\u4efd\u6709\u9650\u516c\u53f8",
    "\u6709\u9650\u8d23\u4efb\u516c\u53f8",
    "\u6709\u9650\u516c\u53f8",
    "\u5408\u4f19\u4f01\u4e1a\uff08\u6709\u9650\u5408\u4f19\uff09",
    "\u5408\u4f19\u4f01\u4e1a(\u6709\u9650\u5408\u4f19)",
    "\u5408\u4f19\u4f01\u4e1a",
    "\u6709\u9650\u5408\u4f19",
    "\u57fa\u91d1",
    "\u96c6\u56e2",
    "\u4e2d\u5fc3",
    "\u516c\u53f8",
)
GENERIC_ORG_NAMES = {
    "\u76ee\u6807\u516c\u53f8",
    "\u96c6\u56e2\u516c\u53f8",
    "\u516c\u53f8",
    "\u672c\u516c\u53f8",
    "\u5404\u96c6\u56e2\u516c\u53f8",
    "\u6295\u8d44\u4eba",
    "\u80a1\u4e1c",
}

CHINESE_NUMERAL_ROMAN = {
    "\u4e00": "I",
    "\u4e8c": "II",
    "\u4e09": "III",
    "\u56db": "IV",
    "\u4e94": "V",
    "\u516d": "VI",
    "\u4e03": "VII",
    "\u516b": "VIII",
    "\u4e5d": "IX",
    "\u5341": "X",
}
ARABIC_ROMAN = {str(index): roman for index, roman in enumerate(["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]) if index}
ROMAN_ARABIC = {roman: arabic for arabic, roman in ARABIC_ROMAN.items()}
APPENDIX_SOURCE_RE = re.compile(r"(\u9644\u5f55|\u9644\u4ef6|\u9644\u8868)([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+)(?:[\uff08(]\s*([A-Za-z\uff21-\uff3a])\s*[\uff09)])?")
SOURCE_FALLBACK_TARGETS = {
    "\u5b9a\u4e49": ["Definitions"],
    "\u672c\u6b21\u4ea4\u6613\u5b89\u6392": ["Transaction Arrangement", "Transaction Arrangements"],
    "\u4ea4\u5272": ["Closing"],
    "\u4ea4\u5272\u540e\u4e49\u52a1": ["Post-Closing Obligations", "Post-closing Obligations"],
    "\u58f0\u660e\u548c\u4fdd\u8bc1": ["Representations and Warranties"],
}

CJK_PUNCT_TRANSLATION = str.maketrans(
    {
        "\uff0c": ",",
        "\u3002": ".",
        "\uff1b": ";",
        "\uff1a": ":",
        "\uff01": "!",
        "\uff1f": "?",
        "\uff08": "(",
        "\uff09": ")",
        "\u201c": '"',
        "\u201d": '"',
        "\u2018": "'",
        "\u2019": "'",
        "\u3001": ",",
        "\u300a": '"',
        "\u300b": '"',
    }
)

BLOCK_RE = re.compile(r"<!--\s*BLOCK:(\d+)\s*-->\s*\n(.*?)(?=\n<!--\s*BLOCK:\d+\s*-->\s*\n|\Z)", re.DOTALL)


@dataclass
class FormatSpan:
    span_id: str
    block_id: int
    source_text: str
    style_label: str
    bold: bool
    italic: bool
    underline: object
    highlight_color: object


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
W_VAL = f"{W_NS}val"
OFF_VALUES = {"0", "false", "off", "none"}


def has_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text or ""))


def compact_text(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if len(text) <= limit:
        return text
    half = max(1, limit // 2)
    return text[:half] + "\n...[middle omitted]...\n" + text[-half:]


def progress_ratio(done: int | float, total: int | float) -> float:
    try:
        total_value = float(total)
        if total_value <= 0:
            return 0.0
        return max(0.0, min(1.0, float(done) / total_value))
    except (TypeError, ValueError, ZeroDivisionError):
        return 0.0


def report_progress(progress, value: int | float, current: str) -> None:
    progress(max(0, min(99, int(round(value)))), PROGRESS_TOTAL, current)


def report_complete(progress, current: str) -> None:
    progress(PROGRESS_TOTAL, PROGRESS_TOTAL, current)


def make_phase_progress(progress, start: int | float, end: int | float):
    def phase(done: int | float, total: int | float, current: str) -> None:
        ratio = progress_ratio(done, total)
        report_progress(progress, start + (end - start) * ratio, current)

    return phase


def report_phase(progress, start: int | float, end: int | float, done: int | float, total: int | float, current: str) -> None:
    make_phase_progress(progress, start, end)(done, total, current)


class TranslationCancelled(Exception):
    pass


_CANCEL_CONTEXT = threading.local()


def set_active_cancel_event(cancel_event: threading.Event | None) -> None:
    _CANCEL_CONTEXT.event = cancel_event


def get_active_cancel_event() -> threading.Event | None:
    return getattr(_CANCEL_CONTEXT, "event", None)


def raise_if_cancelled() -> None:
    cancel_event = get_active_cancel_event()
    if cancel_event and cancel_event.is_set():
        raise TranslationCancelled()


def run_cancellable_api_call(callable_obj):
    cancel_event = get_active_cancel_event()
    if not cancel_event:
        return callable_obj()

    result = {}

    def worker():
        try:
            result["value"] = callable_obj()
        except BaseException as exc:
            result["error"] = exc

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()
    while thread.is_alive():
        if cancel_event.is_set():
            raise TranslationCancelled()
        thread.join(0.2)
    if "error" in result:
        raise result["error"]
    return result.get("value")


def legal_rag_for_text(text: str) -> str:
    retrieved = []
    text = text or ""
    for trigger, note in LEGAL_RAG_TERMS:
        if trigger in text and note not in retrieved:
            retrieved.append(note)
    if not retrieved:
        return CAPITAL_MARKETS_LEGAL_RAG
    return CAPITAL_MARKETS_LEGAL_RAG + "\n\nRetrieved terms for this chunk:\n- " + "\n- ".join(retrieved)


def clean_company_source_name(name: str) -> str:
    name = re.sub(r"\s+", "", name or "")
    name = name.strip(" \t\r\n,.;:!?()[]{}<>\"'")
    name = name.strip("\u3001\uff0c\u3002\uff1b\uff1a\uff08\uff09\u3010\u3011\u300a\u300b\u201c\u201d\u2018\u2019")
    return name


def is_likely_chinese_org_name(name: str) -> bool:
    name = clean_company_source_name(name)
    if not name or name in GENERIC_ORG_NAMES:
        return False
    if len(name) < 4 or not has_cjk(name):
        return False
    if any(word in name for word in ("\u672c\u534f\u8bae", "\u672c\u6b21\u4ea4\u6613", "\u8463\u4e8b\u4f1a", "\u5de5\u5546\u884c\u653f")):
        return False
    return any(suffix in name for suffix in ORG_NAME_SUFFIXES)


ORG_CANDIDATE_RE = re.compile(
    r"[\u4e00-\u9fff][\u4e00-\u9fffA-Za-z0-9\uff08\uff09()·\-\u00b7]{1,70}?"
    r"(?:\u80a1\u4efd\u6709\u9650\u516c\u53f8|\u6709\u9650\u8d23\u4efb\u516c\u53f8|\u6709\u9650\u516c\u53f8|"
    r"\u5408\u4f19\u4f01\u4e1a\uff08\u6709\u9650\u5408\u4f19\uff09|\u5408\u4f19\u4f01\u4e1a\(\u6709\u9650\u5408\u4f19\)|"
    r"\u5408\u4f19\u4f01\u4e1a|\u6709\u9650\u5408\u4f19|\u57fa\u91d1|\u96c6\u56e2|\u4e2d\u5fc3|\u516c\u53f8)"
)


def extract_company_name_candidates(text: str, limit: int = 250) -> list[str]:
    names = []
    seen = set()
    for match in ORG_CANDIDATE_RE.finditer(text or ""):
        name = clean_company_source_name(match.group(0))
        if not is_likely_chinese_org_name(name) or name in seen:
            continue
        seen.add(name)
        names.append(name)
        if len(names) >= limit:
            break
    return names


def clean_company_english_name(name: str, source_name: str) -> str:
    name = normalize_legal_english(name or "")
    source_name = clean_company_source_name(source_name)
    if source_name:
        name = name.replace(source_name, "")
        name = re.sub(r"\(\s*\)", "", name)
    name = re.sub(r"[\u3400-\u9fff]+", "", name)
    name = re.sub(r"\s+", " ", name)
    return name.strip(" -_,;:()[]{}")


def make_company_required_text(source_name: str, english_name: str) -> str:
    source_name = clean_company_source_name(source_name)
    english_name = clean_company_english_name(english_name, source_name)
    if not source_name or not english_name:
        return ""
    return f"{english_name} ({source_name})"


def company_glossary_to_prompt(entries: list[dict]) -> str:
    if not entries:
        return "(none)"
    lines = [
        "Mandatory company/institution name glossary. Whenever the Chinese source name appears in a block, the English translation must include the required rendering exactly, unless the source itself already supplies a different official English name:"
    ]
    for entry in entries[:300]:
        source_name = entry.get("source_name", "")
        required_text = entry.get("required_text", "")
        if source_name and required_text:
            lines.append(f"- {source_name} => {required_text}")
    return "\n".join(lines)


def valid_english_font(font_name: str) -> str:
    return font_name if font_name in ENGLISH_FONT_OPTIONS else DEFAULT_ENGLISH_FONT


def valid_chinese_font(font_name: str) -> str:
    return font_name if font_name in CHINESE_FONT_OPTIONS else DEFAULT_CHINESE_FONT


def build_font_instruction(english_font: str, chinese_font: str) -> str:
    english_font = valid_english_font(english_font)
    chinese_font = valid_chinese_font(chinese_font)
    return (
        f"Font preference for the exported Word document: use {english_font} for English Latin letters, "
        f"use {DIGIT_FONT} for all Arabic numerals/digits. The final English translation should not retain Chinese/CJK "
        f"characters; {chinese_font} is only a fallback font if a source-required symbol or placeholder forces non-Latin text. "
        "Preserve legal meaning even when applying these font preferences."
    )


def normalize_english_punctuation(text: str) -> str:
    text = (text or "").translate(CJK_PUNCT_TRANSLATION)
    text = re.sub(r"\s+([,.;:!?\]\)])", r"\1", text)
    text = re.sub(r"([\[\(])\s+", r"\1", text)
    text = re.sub(r"([,;:!?])(?=[A-Za-z\"'])", r"\1 ", text)
    text = re.sub(r"(?<!\d)(\.)(?=[A-Za-z\"'])", r"\1 ", text)
    text = re.sub(r"(?<=\d),\s+(?=\d{3}\b)", ",", text)
    text = re.sub(r"(?<=\d)\.\s+(?=\d)", ".", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def strip_markdown_emphasis(text: str) -> str:
    text = text or ""
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\*\*([^*\n]+?)\*\*", r"\1", text)
        text = re.sub(r"(?<!\*)\*([^*\n]+?)\*(?!\*)", r"\1", text)
        text = re.sub(r"`([^`\n]+?)`", r"\1", text)
    return text


FULL_RMB_WANYUAN_RE = re.compile(r"^\s*\u4eba\u6c11\u5e01\s*([\u3010\[])?\s*([0-9][0-9,，]*(?:\.\d+)?)\s*([\u3011\]])?\s*\u4e07\u5143\s*$")
FULL_CHINESE_DATE_RE = re.compile(r"^\s*(\d{4})\s*\u5e74\s*(\d{1,2}|\u3010\u3011|_+)\s*\u6708\s*(\d{1,2}|\u3010\u3011|_+)\s*\u65e5\s*$")
SOURCE_BRACKET_RE = re.compile(r"\u3010([^\u3011]*)\u3011")
MONTH_NAMES = {
    1: "January",
    2: "February",
    3: "March",
    4: "April",
    5: "May",
    6: "June",
    7: "July",
    8: "August",
    9: "September",
    10: "October",
    11: "November",
    12: "December",
}


def format_rmb_yuan_from_wanyuan(number_text: str) -> str | None:
    cleaned = (number_text or "").replace(",", "").replace("，", "").strip()
    try:
        amount = (Decimal(cleaned) * Decimal("10000")).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError):
        return None
    return f"{int(amount):,}"


def deterministic_amount_translation(source_text: str) -> str | None:
    match = FULL_RMB_WANYUAN_RE.fullmatch(source_text or "")
    if not match:
        return None
    amount = format_rmb_yuan_from_wanyuan(match.group(2))
    if not amount:
        return None
    if match.group(1) or match.group(3):
        return f"RMB \u3010{amount}\u3011"
    return f"RMB {amount}"


def date_token_to_english(token: str, is_month: bool) -> str:
    token = (token or "").strip()
    if token == "\u3010\u3011" or re.fullmatch(r"_+", token):
        return token
    if not token.isdigit():
        return token
    value = int(token)
    if is_month and value in MONTH_NAMES:
        return MONTH_NAMES[value]
    return str(value)


def deterministic_date_translation(source_text: str) -> str | None:
    match = FULL_CHINESE_DATE_RE.fullmatch(source_text or "")
    if not match:
        return None
    year, month_token, day_token = match.groups()
    month = date_token_to_english(month_token, True)
    day = date_token_to_english(day_token, False)
    return f"{month} {day}, {year}"


def preserve_source_bracket_tokens(source_text: str, translation: str) -> str:
    result = translation or ""
    for content in SOURCE_BRACKET_RE.findall(source_text or ""):
        target = f"\u3010{content}\u3011"
        if target in result:
            continue
        if content and has_cjk(content):
            continue
        if content:
            result = re.sub(r"\[\s*" + re.escape(content) + r"\s*\]", target, result, count=1)
        else:
            result = re.sub(r"\[\s*\]", target, result, count=1)
            result = re.sub(r"_+", target, result, count=1)
    return result


def normalize_translation_against_source(source_text: str, translation: str) -> str:
    deterministic = deterministic_amount_translation(source_text)
    if deterministic:
        return deterministic
    deterministic = deterministic_date_translation(source_text)
    if deterministic:
        return deterministic
    return preserve_source_bracket_tokens(source_text, normalize_legal_english(translation))


def normalize_roman_token(token: str) -> str:
    token = (token or "").strip().upper()
    if token in CHINESE_NUMERAL_ROMAN:
        return CHINESE_NUMERAL_ROMAN[token]
    if token in ARABIC_ROMAN:
        return ARABIC_ROMAN[token]
    if token in ROMAN_ARABIC:
        return token
    return token


def normalize_legal_english(text: str) -> str:
    text = strip_markdown_emphasis(text)
    text = normalize_english_punctuation(text)

    def appendix_repl(match):
        roman = normalize_roman_token(match.group(1))
        suffix = f"({match.group(2).upper()})" if match.group(2) else ""
        return f"Appendix {roman}{suffix}"

    text = re.sub(r"\bAppendix\s+([1-9]|10|I|II|III|IV|V|VI|VII|VIII|IX|X)(?!\.\d)(?:\s*\(\s*([A-Za-z])\s*\))?", appendix_repl, text, flags=re.IGNORECASE)
    text = re.sub(r"\bSchedule\s+([1-9]|10|I|II|III|IV|V|VI|VII|VIII|IX|X)\s*\(?\s*Disclosure Schedule\s*\)?", lambda m: f"Appendix {normalize_roman_token(m.group(1))} (Disclosure Schedule)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bSchedule\s+([1-9]|10|I|II|III|IV|V|VI|VII|VIII|IX|X)\s+of\s+the\s+Disclosure\s+Schedule\b", lambda m: f"Appendix {normalize_roman_token(m.group(1))} (Disclosure Schedule)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bLong\s+Stop\s+Date\b", "Longstop Date", text, flags=re.IGNORECASE)
    return text


def source_span_fallback_targets(source_text: str) -> list[str]:
    source_text = source_text or ""
    direct_targets = SOURCE_FALLBACK_TARGETS.get(source_text.strip(), [])
    if direct_targets:
        return direct_targets
    match = APPENDIX_SOURCE_RE.search(source_text or "")
    if not match:
        return []
    label_map = {"\u9644\u5f55": "Appendix", "\u9644\u4ef6": "Annex", "\u9644\u8868": "Schedule"}
    label = label_map.get(match.group(1))
    roman = normalize_roman_token(match.group(2))
    if not label or not roman:
        return []
    suffix = ""
    if match.group(3):
        suffix = f"({match.group(3).upper()})"
    targets = [f"{label} {roman}{suffix}"]
    arabic = ROMAN_ARABIC.get(roman)
    if arabic:
        targets.append(f"{label} {arabic}{suffix}")
    return targets


def load_config() -> dict:
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def save_config(provider: str, api_key: str, base_url: str, model: str, english_font: str, chinese_font: str) -> None:
    CONFIG_PATH.write_text(
        json.dumps(
            {
                "provider": provider,
                "api_key": api_key,
                "base_url": base_url,
                "model": model,
                "english_font": english_font,
                "chinese_font": chinese_font,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def build_client(api_key: str, base_url: str) -> OpenAI:
    kwargs = {"api_key": api_key, "timeout": 180}
    if base_url.strip():
        kwargs["base_url"] = base_url.strip()
    return OpenAI(**kwargs)


def parse_json_object(raw: str) -> dict:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r"```$", "", raw).strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("模型没有返回 JSON 对象。")
    return json.loads(raw[start : end + 1])


def chat_json(client: OpenAI, provider: str, model: str, messages: list[dict], retries: int = 3) -> dict:
    last_error = None
    for attempt in range(retries):
        raise_if_cancelled()
        response_format_options = [True, False] if provider == "DeepSeek" else [True]
        for use_response_format in response_format_options:
            try:
                raise_if_cancelled()
                kwargs = {"model": model, "messages": messages, "temperature": 0.05}
                if use_response_format:
                    kwargs["response_format"] = {"type": "json_object"}
                response = run_cancellable_api_call(lambda: client.chat.completions.create(**kwargs))
                raise_if_cancelled()
                return parse_json_object(response.choices[0].message.content)
            except TranslationCancelled:
                raise
            except Exception as exc:
                last_error = exc
        if attempt + 1 < retries:
            sleep_until = time.time() + 2 + attempt * 3
            while time.time() < sleep_until:
                raise_if_cancelled()
                time.sleep(0.2)
    raise last_error


def make_company_name_glossary(client: OpenAI, provider: str, model: str, source_text: str, log) -> list[dict]:
    candidates = extract_company_name_candidates(source_text)
    if not candidates:
        return []
    log(f"{APP_VERSION} extracting company/institution name glossary: {len(candidates)} candidates.")
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Extract a mandatory company/institution name glossary for a Chinese capital-markets legal translation.\n"
                "Return JSON exactly like {\"entries\":[{\"source_name\":\"中文全称\",\"english_name\":\"Official or best legal English name\"}]}.\n\n"
                "Rules:\n"
                "1. Include complete Chinese legal entity, fund, partnership, shareholder, investor, institution, and group company names.\n"
                "2. Exclude generic legal terms such as 目标公司, 集团公司, 投资人, 股东, 公司, 本公司.\n"
                "3. Prefer the official English name that would appear on Qichacha or in corporate registration records when known.\n"
                "4. If the official English name is unknown, provide a conservative legal English rendering or pinyin-style legal name, but do not omit the entry.\n"
                "5. english_name must not include the Chinese source name or Chinese parentheses; the program will append the Chinese name.\n\n"
                f"SOURCE CANDIDATES:\n{json.dumps(candidates, ensure_ascii=False)}\n\n"
                f"SOURCE EXCERPT:\n{compact_text(source_text, 70000)}"
            ),
        },
    ]
    try:
        data = chat_json(client, provider, model, messages, retries=2)
    except TranslationCancelled:
        raise
    except Exception as exc:
        log(f"{APP_VERSION} company glossary extraction failed; continuing without mandatory name glossary. Reason: {exc}")
        return []

    entries = []
    seen = set()
    for item in data.get("entries", []) or []:
        source_name = clean_company_source_name(str(item.get("source_name", "")))
        english_name = str(item.get("english_name", "")).strip()
        if not is_likely_chinese_org_name(source_name) or source_name in seen:
            continue
        if source_name not in source_text:
            continue
        required_text = make_company_required_text(source_name, english_name)
        if not required_text:
            continue
        seen.add(source_name)
        entries.append(
            {
                "source_name": source_name,
                "english_name": clean_company_english_name(english_name, source_name),
                "required_text": required_text,
            }
        )
    if entries:
        log(f"{APP_VERSION} company/institution glossary ready: {len(entries)} mandatory names.")
    else:
        log(f"{APP_VERSION} company/institution glossary returned no usable entries.")
    return entries


def iter_table_paragraphs(table, seen):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_cell_paragraphs(cell, seen)


def iter_cell_paragraphs(cell, seen):
    for paragraph in cell.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in cell.tables:
        yield from iter_table_paragraphs(table, seen)


def iter_part_paragraphs(part, seen):
    for paragraph in part.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in part.tables:
        yield from iter_table_paragraphs(table, seen)


def iter_document_paragraphs(doc: Document):
    seen = set()
    yield from iter_part_paragraphs(doc, seen)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            yield from iter_part_paragraphs(part, seen)


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def xml_paragraph_text(xml_paragraph) -> str:
    return "".join(node.text or "" for node in xml_paragraph.xpath(".//w:t"))


def rewrite_xml_paragraph_text(xml_paragraph, text: str) -> None:
    text_nodes = xml_paragraph.xpath(".//w:t")
    if not text_nodes:
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        run.append(text_node)
        xml_paragraph.append(run)
        text_nodes = [text_node]
    text_nodes[0].text = text
    text_nodes[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for node in text_nodes[1:]:
        node.text = ""


def paragraph_meta(paragraph) -> dict:
    fmt = paragraph.paragraph_format
    return {
        "style": paragraph.style.name if paragraph.style else "",
        "alignment": str(paragraph.alignment) if paragraph.alignment else "",
        "left_indent": str(fmt.left_indent) if fmt.left_indent else "",
        "first_line_indent": str(fmt.first_line_indent) if fmt.first_line_indent else "",
    }


def markdown_prefix(meta: dict, text: str) -> str:
    style = (meta.get("style") or "").lower()
    stripped = text.strip()
    if "heading 1" in style or style in ("title", "标题 1"):
        return "# "
    if "heading 2" in style or style == "标题 2":
        return "## "
    if "heading 3" in style or style == "标题 3":
        return "### "
    if re.match(r"^第[一二三四五六七八九十百零〇]+条", stripped):
        return "## "
    if re.match(r"^\d+(\.\d+)*\s+", stripped):
        return ""
    return ""


def block_to_markdown(block: dict) -> str:
    meta = block["meta"]
    prefix = markdown_prefix(meta, block["text"])
    meta_json = json.dumps(meta, ensure_ascii=False)
    return f"<!-- BLOCK:{block['id']} -->\n{prefix}{block['text']}\n<!-- META:{meta_json} -->"


def parse_translated_markdown(markdown: str) -> dict[int, str]:
    results = {}
    for match in BLOCK_RE.finditer(markdown.strip()):
        block_id = int(match.group(1))
        body = match.group(2).strip()
        body = re.sub(r"\n<!--\s*META:.*?-->\s*$", "", body, flags=re.DOTALL).strip()
        body = re.sub(r"^#{1,6}\s+", "", body).strip()
        results[block_id] = normalize_legal_english(body)
    return results


def build_translated_blocks_markdown(blocks: list[dict], translated_blocks: dict[int, str]) -> str:
    parts = []
    for block in blocks:
        block_id = int(block["id"])
        if block_id not in translated_blocks:
            continue
        meta_json = json.dumps(block.get("meta", {}), ensure_ascii=False)
        parts.append(f"<!-- BLOCK:{block_id} -->\n{translated_blocks[block_id]}\n<!-- META:{meta_json} -->")
    return "\n\n".join(parts)


def xml_has_off_property(element, tag_name: str) -> bool:
    if element is None:
        return False
    for child in element.iter(f"{W_NS}{tag_name}"):
        value = child.get(W_VAL)
        if value is not None and value.lower() in OFF_VALUES:
            return True
    return False


def xml_has_on_property(element, tag_name: str) -> bool:
    if element is None:
        return False
    for child in element.iter(f"{W_NS}{tag_name}"):
        value = child.get(W_VAL)
        if value is None or value.lower() not in OFF_VALUES:
            return True
    return False


def style_format_defaults(paragraph) -> tuple:
    style = getattr(paragraph, "style", None)
    font = getattr(style, "font", None)
    if font is None:
        return (False, False, None, None)
    return (bool(font.bold), bool(font.italic), font.underline, font.highlight_color)


def run_style_tuple(run, paragraph=None) -> tuple:
    direct = (bool(run.bold), bool(run.italic), run.underline, run.font.highlight_color)
    if paragraph is None:
        return direct
    style_bold, style_italic, style_underline, style_highlight = style_format_defaults(paragraph)
    rpr = run._r.rPr

    bold = True if run.bold is True or xml_has_on_property(rpr, "b") else False
    if not bold and run.bold is not False and not xml_has_off_property(rpr, "b") and style_bold:
        bold = True

    italic = True if run.italic is True or xml_has_on_property(rpr, "i") else False
    if not italic and run.italic is not False and not xml_has_off_property(rpr, "i") and style_italic:
        italic = True

    underline = run.underline
    if not underline and run.underline is not False and not xml_has_off_property(rpr, "u"):
        underline = style_underline if style_underline else (True if xml_has_on_property(rpr, "u") else None)

    highlight = run.font.highlight_color
    if not highlight and not xml_has_off_property(rpr, "highlight"):
        highlight = style_highlight
    return (bold, italic, underline, highlight)


def is_formatted_run(run, paragraph=None) -> bool:
    bold, italic, underline, highlight = run_style_tuple(run, paragraph)
    return bool((bold or italic or underline or highlight) and run.text and run.text.strip())


def style_label(style: tuple) -> str:
    bold, italic, underline, highlight = style
    labels = []
    if bold:
        labels.append("bold")
    if italic:
        labels.append("italic")
    if underline:
        labels.append("underline")
    if highlight:
        labels.append(f"highlight:{highlight}")
    return "+".join(labels)


def extract_format_spans(paragraph, block_id: int) -> list[FormatSpan]:
    spans = []
    current_text = ""
    current_style = None
    span_number = 0

    def flush():
        nonlocal current_text, current_style, span_number
        if current_text.strip() and current_style:
            bold, italic, underline, highlight = current_style
            spans.append(
                FormatSpan(
                    span_id=f"b{block_id}_s{span_number}",
                    block_id=block_id,
                    source_text=current_text,
                    style_label=style_label(current_style),
                    bold=bold,
                    italic=italic,
                    underline=copy(underline),
                    highlight_color=copy(highlight),
                )
            )
            span_number += 1
        current_text = ""
        current_style = None

    for run in paragraph.runs:
        if is_formatted_run(run, paragraph):
            style = run_style_tuple(run, paragraph)
            if style == current_style:
                current_text += run.text
            else:
                flush()
                current_style = style
                current_text = run.text
        else:
            flush()
    flush()
    return spans


def build_chunks(blocks: list[dict], max_chars: int = 35000, max_blocks: int = 45) -> list[list[dict]]:
    chunks = []
    current = []
    current_chars = 0
    for block in blocks:
        block_chars = len(block["markdown"]) + sum(len(span["source_text"]) for span in block["format_spans"])
        if current and (current_chars + block_chars > max_chars or len(current) >= max_blocks):
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(block)
        current_chars += block_chars
    if current:
        chunks.append(current)
    return chunks


def make_translation_memory(client: OpenAI, provider: str, model: str, source_markdown: str, log, font_instruction: str, company_glossary: str) -> str:
    log("正在基于 Markdown 结构提取术语表和翻译规则...")
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Create a concise bilingual translation memory for this Chinese legal contract represented as Markdown blocks.\n"
                "Return JSON {\"memory\":\"...\"}. Include defined terms, parties, recurring phrases, bracket placeholders, and legal drafting preferences.\n\n"
                f"FONT INSTRUCTIONS:\n{font_instruction}\n\n"
                f"COMPANY/INSTITUTION NAME GLOSSARY:\n{company_glossary}\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(source_markdown)}\n\n"
                f"MARKDOWN:\n{compact_text(source_markdown, 60000)}"
            ),
        },
    ]
    try:
        return str(chat_json(client, provider, model, messages, retries=2).get("memory", "")).strip()
    except TranslationCancelled:
        raise
    except Exception as exc:
        log(f"术语记忆提取失败，将继续翻译。原因：{exc}")
        return ""


def translate_markdown_chunk(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    chunk_markdown: str,
    before: str,
    after: str,
    font_instruction: str,
    company_glossary: str,
) -> str:
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate this Markdown-structured Chinese legal contract chunk into precise formal legal English.\n"
                "Return JSON {\"markdown\":\"...\"}.\n"
                "Rules:\n"
                "1. Preserve every <!-- BLOCK:n --> marker exactly.\n"
                "2. Preserve Markdown heading/list/paragraph structure where reasonable.\n"
                "3. Translate the visible Chinese text into English using the whole chunk context.\n"
                "4. Do not translate or remove <!-- META:... --> comments.\n"
                "5. Do not leave any Chinese/CJK characters in visible translated text. Do not use English Name (Chinese Name) formatting.\n"
                "6. Translate or romanize Chinese company, fund, person, address, exhibit, appendix, and brand/trademark names into English-only text.\n\n"
                "Workflow reminder:\n"
                "A. Markdown carries clause hierarchy, indentation/list cues, and some visible formatting signals, but it is not the complete Word run structure.\n"
                "B. Translate under the Markdown structure first so the English keeps context and legal coherence.\n"
                "C. Keep each BLOCK independent after translation because run-level formatting will be checked and applied block by block.\n\n"
                f"FONT INSTRUCTIONS:\n{font_instruction}\n\n"
                f"COMPANY/INSTITUTION NAME GLOSSARY:\n{company_glossary}\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(chunk_markdown)}\n\n"
                f"TRANSLATION MEMORY:\n{memory or '(none)'}\n\n"
                f"CONTEXT BEFORE:\n{compact_text(before, 3000)}\n\n"
                f"CONTEXT AFTER:\n{compact_text(after, 3000)}\n\n"
                f"CHUNK MARKDOWN:\n{chunk_markdown}"
            ),
        },
    ]
    return str(chat_json(client, provider, model, messages).get("markdown", "")).strip()


def cjk_count(text: str) -> int:
    return len(re.findall(r"[\u3400-\u9fff]", text or ""))


def needs_translation_repair(source_text: str, translation: str) -> bool:
    translation = translation or ""
    if not translation.strip():
        return True
    count = cjk_count(translation)
    if not count:
        return False
    if STRICT_NO_CJK_OUTPUT:
        return True
    compact_len = max(1, len(re.sub(r"\s+", "", translation)))
    source_count = max(1, cjk_count(source_text))
    ascii_letters = len(re.findall(r"[A-Za-z]", translation))
    if count >= 4 and count / compact_len >= 0.5 and ascii_letters < 10:
        return True
    return (count >= 20 and count / compact_len >= 0.08) or count >= max(30, source_count // 4)


def build_repair_batches(blocks: list[dict], max_chars: int = 12000, max_blocks: int = 8) -> list[list[dict]]:
    batches = []
    current = []
    current_chars = 0
    for block in blocks:
        size = len(block.get("text", "")) + 200
        if current and (current_chars + size > max_chars or len(current) >= max_blocks):
            batches.append(current)
            current = []
            current_chars = 0
        current.append(block)
        current_chars += size
    if current:
        batches.append(current)
    return batches


def repair_translation_batch(client: OpenAI, provider: str, model: str, memory: str, batch: list[dict], font_instruction: str, company_glossary: str) -> dict[int, str]:
    payload = [{"id": block["id"], "source_text": block["text"]} for block in batch]
    source_text = "\n".join(item["source_text"] for item in payload)
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Repair missing or incomplete translations for these Chinese legal contract blocks.\n"
                "Return JSON exactly like {\"translations\":[{\"id\":123,\"translation\":\"...\"}]}.\n"
                "Rules:\n"
                "1. Return one translation for every input id.\n"
                "2. Translate all Chinese legal prose into formal legal English.\n"
                "3. Do not leave any Chinese/CJK characters in the translation.\n"
                "4. Do not use English Name (Chinese Name) formatting; render Chinese company, fund, person, address, and trademark/brand names in English-only form by using the official English name where known or a reasonable romanized/legal-English rendering.\n"
                "5. Preserve numbers, defined terms, bracket symbols, blanks, placeholders, clause references, and round names such as B5 Round / Pre-IPO Round. If Chinese appears inside 【】, translate the bracket content but keep the 【】 symbols.\n\n"
                f"FONT INSTRUCTIONS:\n{font_instruction}\n\n"
                f"COMPANY/INSTITUTION NAME GLOSSARY:\n{company_glossary}\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(source_text)}\n\n"
                f"TRANSLATION MEMORY:\n{memory or '(none)'}\n\n"
                f"BLOCKS JSON:\n{json.dumps({'blocks': payload}, ensure_ascii=False)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages, retries=2)
    repaired = {}
    for item in data.get("translations", []) or []:
        try:
            block_id = int(item.get("id"))
        except (TypeError, ValueError):
            continue
        translation = str(item.get("translation", "")).strip()
        if translation:
            repaired[block_id] = translation
    return repaired


def repair_chunk_translations(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    chunk: list[dict],
    translations: dict[int, str],
    log,
    progress,
    completed: int,
    total: int,
    font_instruction: str,
    company_glossary: str,
) -> set[int]:
    block_by_id = {int(block["id"]): block for block in chunk}
    bad_blocks = []
    for block in chunk:
        block_id = int(block["id"])
        translation = translations.get(block_id, "")
        if needs_translation_repair(block["text"], translation):
            bad_blocks.append(block)
    if not bad_blocks:
        return set()

    repaired_ids: set[int] = set()
    ids = ", ".join(str(block["id"]) for block in bad_blocks[:12])
    more = "..." if len(bad_blocks) > 12 else ""
    log(f"自检查发现 {len(bad_blocks)} 个 block 漏翻或中文残留，自动拆分重翻：{ids}{more}")
    batches = build_repair_batches(bad_blocks)
    for index, batch in enumerate(batches, start=1):
        progress(completed, total, f"自修复漏翻/中文残留 {index}/{len(batches)}")
        try:
            repaired = repair_translation_batch(client, provider, model, memory, batch, font_instruction, company_glossary)
        except TranslationCancelled:
            raise
        except Exception as exc:
            log(f"批量自修复失败，改为逐段修复。原因：{exc}")
            repaired = {}
        for block_id, translation in repaired.items():
            block = block_by_id.get(block_id)
            if not block:
                continue
            translation = normalize_translation_against_source(block["text"], translation)
            if translation and not needs_translation_repair(block["text"], translation):
                translations[block_id] = translation
                repaired_ids.add(block_id)
        for block in batch:
            block_id = int(block["id"])
            if block_id in repaired_ids and not needs_translation_repair(block["text"], translations.get(block_id, "")):
                continue
            try:
                single = repair_translation_batch(client, provider, model, memory, [block], font_instruction, company_glossary)
            except TranslationCancelled:
                raise
            except Exception as exc:
                log(f"block {block_id} 单段自修复失败：{exc}")
                continue
            translation = normalize_translation_against_source(block["text"], single.get(block_id, ""))
            if translation:
                translations[block_id] = translation
                repaired_ids.add(block_id)
    return repaired_ids


def repair_all_translations(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    blocks: list[dict],
    translations: dict[int, str],
    log,
    progress,
    font_instruction: str,
    company_glossary: str,
    max_rounds: int = 5,
) -> set[int]:
    repaired_ids: set[int] = set()
    block_by_id = {int(block["id"]): block for block in blocks}
    total = len(blocks)
    for round_index in range(1, max_rounds + 1):
        bad_blocks = [
            block
            for block in blocks
            if needs_translation_repair(block["text"], translations.get(int(block["id"]), ""))
        ]
        if not bad_blocks:
            log(f"{APP_VERSION} 全文自检查通过：没有发现明显漏翻或大段中文残留。")
            return repaired_ids

        ids = ", ".join(str(block["id"]) for block in bad_blocks[:20])
        more = "..." if len(bad_blocks) > 20 else ""
        log(f"{APP_VERSION} 全文自检查第 {round_index} 轮：发现 {len(bad_blocks)} 个 block 需要修复：{ids}{more}")
        changed = False
        batches = build_repair_batches(bad_blocks)
        for batch_index, batch in enumerate(batches, start=1):
            progress(total - len(bad_blocks), total, f"全文自修复第 {round_index} 轮 {batch_index}/{len(batches)}")
            try:
                repaired = repair_translation_batch(client, provider, model, memory, batch, font_instruction, company_glossary)
            except TranslationCancelled:
                raise
            except Exception as exc:
                log(f"批量自修复失败，改为逐段修复。原因：{exc}")
                repaired = {}
            for block_id, translation in repaired.items():
                block = block_by_id.get(block_id)
                if not block:
                    continue
                translation = normalize_translation_against_source(block["text"], translation)
                if translation and translation != translations.get(block_id, ""):
                    translations[block_id] = translation
                    repaired_ids.add(block_id)
                    changed = True

            for block in batch:
                block_id = int(block["id"])
                if not needs_translation_repair(block["text"], translations.get(block_id, "")):
                    continue
                try:
                    single = repair_translation_batch(client, provider, model, memory, [block], font_instruction, company_glossary)
                except TranslationCancelled:
                    raise
                except Exception as exc:
                    log(f"block {block_id} 单段自修复失败：{exc}")
                    continue
                translation = normalize_translation_against_source(block["text"], single.get(block_id, ""))
                if translation and translation != translations.get(block_id, ""):
                    translations[block_id] = translation
                    repaired_ids.add(block_id)
                    changed = True

        if not changed:
            remaining = [
                str(block["id"])
                for block in blocks
                if needs_translation_repair(block["text"], translations.get(int(block["id"]), ""))
            ]
            log(
                f"{APP_VERSION} 全文自修复本轮没有取得新结果；将继续导出，并在 checklist 标出剩余疑似问题 block："
                f"{', '.join(remaining[:20])}{'...' if len(remaining) > 20 else ''}"
            )
            break
    return repaired_ids


def build_final_audit_batches(blocks: list[dict], translations: dict[int, str], max_chars: int = 50000, max_blocks: int = 80) -> list[list[dict]]:
    batches = []
    current = []
    current_chars = 0
    for block in blocks:
        block_id = int(block["id"])
        size = len(block.get("text", "")) + len(translations.get(block_id, "")) + 300
        if current and (current_chars + size > max_chars or len(current) >= max_blocks):
            batches.append(current)
            current = []
            current_chars = 0
        current.append(block)
        current_chars += size
    if current:
        batches.append(current)
    return batches


def deterministic_company_name_audit_issues(batch: list[dict], translations: dict[int, str], company_entries: list[dict]) -> list[dict]:
    issues = []
    if not company_entries:
        return issues
    for block in batch:
        block_id = int(block["id"])
        source_text = block.get("text", "")
        translation = translations.get(block_id, "")
        for entry in company_entries:
            source_name = entry.get("source_name", "")
            required_text = entry.get("required_text", "")
            if not source_name or source_name not in source_text:
                continue
            if source_name in translation:
                continue
            issues.append(
                {
                    "id": block_id,
                    "reason": f"Company/institution name must follow glossary rendering: {required_text}",
                    "chinese_fragment": source_name,
                }
            )
            break
    return issues


def final_llm_audit_batch(
    client: OpenAI,
    provider: str,
    model: str,
    batch: list[dict],
    translations: dict[int, str],
    company_entries: list[dict],
    company_glossary: str,
) -> list[dict]:
    payload = []
    rag_text_parts = []
    for block in batch:
        block_id = int(block["id"])
        source_text = block.get("text", "")
        translation = translations.get(block_id, "")
        payload.append(
            {
                "id": block_id,
                "source_text": compact_text(source_text, 1600),
                "translation": compact_text(translation, 2600),
            }
        )
        rag_text_parts.append(source_text)
    rag_text = "\n".join(rag_text_parts)
    messages = [
        {"role": "system", "content": FINAL_AUDIT_SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Quickly audit these completed English translation blocks.\n"
                "Return JSON exactly like {\"issues\":[{\"id\":123,\"reason\":\"...\",\"chinese_fragment\":\"...\"}]}.\n"
                "Return an empty issues array if all blocks pass.\n\n"
                "What to flag:\n"
                "- Any Chinese/CJK character remaining in the English translation.\n"
                "- Chinese legal prose, headings, labels, table cells, TOC entries, sentence fragments, or parentheticals.\n"
                "- Chinese company, fund, shareholder, investor, person, trademark, brand, addressee, or address names.\n"
                "- Chinese text appended to or mixed into an otherwise English sentence.\n\n"
                "What not to flag:\n"
                "- The bracket symbols 【】 themselves or blank placeholders such as 【】, as long as the text inside them is not Chinese/CJK.\n"
                "- English-only romanized or official English proper names.\n\n"
                f"COMPANY/INSTITUTION NAME GLOSSARY:\n{company_glossary}\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(rag_text)}\n\n"
                f"BLOCKS JSON:\n{json.dumps({'blocks': payload}, ensure_ascii=False)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages, retries=2)
    batch_ids = {int(block["id"]) for block in batch}
    issues = deterministic_company_name_audit_issues(batch, translations, company_entries)
    issue_ids = {int(issue["id"]) for issue in issues}
    if STRICT_NO_CJK_OUTPUT:
        for block in batch:
            block_id = int(block["id"])
            translation = translations.get(block_id, "")
            if cjk_count(translation) and block_id not in issue_ids:
                fragment = "".join(re.findall(r"[\u3400-\u9fff]+", translation))[:80]
                issues.append(
                    {
                        "id": block_id,
                        "reason": "Strict final-output policy: translation still contains Chinese/CJK characters.",
                        "chinese_fragment": fragment,
                    }
                )
                issue_ids.add(block_id)
    for item in data.get("issues", []) or []:
        try:
            block_id = int(item.get("id"))
        except (TypeError, ValueError):
            continue
        if block_id in batch_ids and block_id not in issue_ids:
            issues.append(
                {
                    "id": block_id,
                    "reason": str(item.get("reason", "")).strip(),
                    "chinese_fragment": str(item.get("chinese_fragment", "")).strip(),
                }
            )
            issue_ids.add(block_id)
    return issues


def final_llm_audit_and_repair_translations(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    blocks: list[dict],
    translations: dict[int, str],
    log,
    progress,
    font_instruction: str,
    company_entries: list[dict],
    company_glossary: str,
    max_rounds: int = 2,
) -> set[int]:
    repaired_ids: set[int] = set()
    block_by_id = {int(block["id"]): block for block in blocks}
    total = len(blocks)
    for round_index in range(1, max_rounds + 1):
        progress(total, total, f"{APP_VERSION} final LLM audit round {round_index}")
        issues = []
        for batch_index, batch in enumerate(build_final_audit_batches(blocks, translations), start=1):
            progress(total, total, f"{APP_VERSION} final LLM audit {batch_index}")
            try:
                issues.extend(final_llm_audit_batch(client, provider, model, batch, translations, company_entries, company_glossary))
            except TranslationCancelled:
                raise
            except Exception as exc:
                log(f"{APP_VERSION} final LLM audit batch failed and was skipped: {exc}")
        if not issues:
            log(f"{APP_VERSION} final LLM audit passed: no untranslated Chinese/CJK legal content found.")
            return repaired_ids

        seen = set()
        bad_blocks = []
        notes = []
        for issue in issues:
            block_id = int(issue["id"])
            if block_id in seen or block_id not in block_by_id:
                continue
            seen.add(block_id)
            bad_blocks.append(block_by_id[block_id])
            fragment = issue.get("chinese_fragment") or issue.get("reason") or ""
            notes.append(f"{block_id}: {compact_text(fragment, 80)}")
        log(
            f"{APP_VERSION} final LLM audit round {round_index}: found {len(bad_blocks)} blocks needing repair: "
            f"{', '.join(notes[:12])}{'...' if len(notes) > 12 else ''}"
        )

        changed = False
        for batch_index, batch in enumerate(build_repair_batches(bad_blocks), start=1):
            progress(total, total, f"{APP_VERSION} final LLM repair {round_index}-{batch_index}")
            try:
                repaired = repair_translation_batch(client, provider, model, memory, batch, font_instruction, company_glossary)
            except TranslationCancelled:
                raise
            except Exception as exc:
                log(f"{APP_VERSION} final LLM repair batch failed, trying single blocks. Reason: {exc}")
                repaired = {}
            for block_id, translation in repaired.items():
                block = block_by_id.get(block_id)
                if not block:
                    continue
                translation = normalize_translation_against_source(block["text"], translation)
                if translation and translation != translations.get(block_id, ""):
                    translations[block_id] = translation
                    repaired_ids.add(block_id)
                    changed = True
            for block in batch:
                block_id = int(block["id"])
                if block_id in repaired_ids and not needs_translation_repair(block["text"], translations.get(block_id, "")):
                    continue
                try:
                    single = repair_translation_batch(client, provider, model, memory, [block], font_instruction, company_glossary)
                except TranslationCancelled:
                    raise
                except Exception as exc:
                    log(f"{APP_VERSION} final LLM single repair failed for block {block_id}: {exc}")
                    continue
                translation = normalize_translation_against_source(block["text"], single.get(block_id, ""))
                if translation and translation != translations.get(block_id, ""):
                    translations[block_id] = translation
                    repaired_ids.add(block_id)
                    changed = True
        if not changed:
            log(f"{APP_VERSION} final LLM audit found issues but repair produced no new translations; exporting with checklist notes.")
            break
    return repaired_ids


def map_format_spans(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    source_blocks: list[dict],
    translated_blocks: dict[int, str],
    font_instruction: str,
) -> list[dict]:
    payload = []
    for block in source_blocks:
        spans = block["format_spans"]
        if spans:
            payload.append(
                {
                    "block_id": block["id"],
                    "source_text": block["text"],
                    "translation": translated_blocks.get(block["id"], ""),
                    "format_spans": spans,
                }
            )
    if not payload:
        return []
    payload_text = "\n".join(f"{item['source_text']}\n{item['translation']}" for item in payload)
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Map each formatted Chinese source span to the exact English substring in the translated block that should receive the same formatting.\n"
                "Return JSON {\"mappings\":[{\"span_id\":\"...\",\"block_id\":1,\"target_text\":\"...\",\"confidence\":\"high|medium|low\",\"note\":\"...\"}]}.\n"
                "Rules:\n"
                "1. Every input span_id must appear once.\n"
                "2. target_text should be an exact contiguous substring of the English translation whenever possible.\n"
                "3. If the concept is split in English, choose the closest legally equivalent contiguous phrase and mark confidence low.\n"
                "4. Do not invent styling; only map the supplied spans.\n\n"
                "Mandatory block-by-block checklist:\n"
                "A. Treat Markdown as already carrying only part of the document structure/formatting.\n"
                "B. For each block, inspect the run-derived format_spans as the checklist of formatting not fully represented by Markdown.\n"
                "C. For every span, record the formatted source content, locate the legally corresponding exact phrase in that block's English translation, and return it as target_text.\n"
                "D. Finish all spans in the current block before moving to the next block.\n"
                "E. If a target is uncertain, still return the closest contiguous legally equivalent English phrase and explain the uncertainty in note.\n\n"
                f"FONT INSTRUCTIONS:\n{font_instruction}\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(payload_text)}\n\n"
                f"TRANSLATION MEMORY:\n{memory or '(none)'}\n\n"
                f"BLOCKS JSON:\n{json.dumps({'blocks': payload}, ensure_ascii=False)}"
            ),
        },
    ]
    return chat_json(client, provider, model, messages).get("mappings", []) or []


def nearest_nonspace(text: str, index: int, direction: int) -> str:
    while 0 <= index < len(text):
        if not text[index].isspace():
            return text[index]
        index += direction
    return ""


def candidate_match_score(text: str, start: int, end: int, exact_case: bool) -> int:
    score = 20 if exact_case else 0
    before = nearest_nonspace(text, start - 1, -1)
    after = nearest_nonspace(text, end, 1)
    if before in {'"', "'", "\u201c", "\u2018"} and after in {'"', "'", "\u201d", "\u2019"}:
        score += 100
    if before in {"(", "["} and after in {")", "]"}:
        score += 20
    if start == 0 or not text[start - 1].isalnum():
        score += 8
    else:
        score -= 50
    if end == len(text) or not text[end : end + 1].isalnum():
        score += 8
    else:
        score -= 50
    prefix = text[max(0, start - 8) : start].lower()
    if '"' in prefix or "\u201c" in prefix or "(" in prefix:
        score += 5
    return score


def find_available_span(text: str, needle: str, occupied: list[tuple[int, int]]) -> tuple[int, int] | None:
    if not needle:
        return None
    candidates = []
    for candidate in [needle, needle.strip(), re.sub(r"\s+", " ", needle).strip(), normalize_legal_english(needle)]:
        if candidate and candidate not in candidates:
            candidates.append(candidate)
    for candidate in candidates:
        if not candidate:
            continue
        matches = []
        start = text.find(candidate)
        while start != -1:
            end = start + len(candidate)
            if all(end <= a or start >= b for a, b in occupied):
                matches.append((candidate_match_score(text, start, end, True), start, end))
            start = text.find(candidate, start + 1)
        lower_text = text.lower()
        lower_candidate = candidate.lower()
        start = lower_text.find(lower_candidate)
        while start != -1:
            end = start + len(candidate)
            if all(end <= a or start >= b for a, b in occupied):
                matches.append((candidate_match_score(text, start, end, False), start, end))
            start = lower_text.find(lower_candidate, start + 1)
        if matches:
            matches.sort(key=lambda item: (-item[0], item[1]))
            _, start, end = matches[0]
            return start, end
    return None


def clear_paragraph_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)


def paragraph_rpr(paragraph):
    ppr = paragraph._p.pPr
    return None if ppr is None else ppr.find(qn("w:rPr"))


def style_font_size(style):
    while style is not None:
        font = getattr(style, "font", None)
        size = getattr(font, "size", None) if font is not None else None
        if size:
            return size
        style = getattr(style, "base_style", None)
    return None


def run_font_size(run):
    size = getattr(run.font, "size", None)
    if size:
        return size
    return style_font_size(getattr(run, "style", None))


def paragraph_font_size(paragraph):
    for run in paragraph.runs:
        if not run.text:
            continue
        size = run_font_size(run)
        if size:
            return size
    return style_font_size(getattr(paragraph, "style", None))


def base_format_overrides(paragraph) -> dict[str, object]:
    overrides = {
        "bold": False,
        "italic": False,
        "underline": False,
        "highlight": False,
        "font_size": paragraph_font_size(paragraph),
    }
    rpr = paragraph_rpr(paragraph)
    if xml_has_off_property(rpr, "b"):
        overrides["bold"] = True
    if xml_has_off_property(rpr, "i"):
        overrides["italic"] = True
    if xml_has_off_property(rpr, "u"):
        overrides["underline"] = True
    if xml_has_off_property(rpr, "highlight"):
        overrides["highlight"] = True

    for run in paragraph.runs:
        if not run.text:
            continue
        if run.bold is False:
            overrides["bold"] = True
        if run.italic is False:
            overrides["italic"] = True
        if run.underline is False:
            overrides["underline"] = True
        if xml_has_off_property(run._r.rPr, "highlight"):
            overrides["highlight"] = True
    return overrides


def set_highlight_none(run) -> None:
    rpr = run._r.get_or_add_rPr()
    for existing in list(rpr.findall(qn("w:highlight"))):
        rpr.remove(existing)
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "none")
    rpr.append(highlight)


def apply_base_format(run, base_format: dict[str, object]) -> None:
    if base_format.get("bold"):
        run.bold = False
    if base_format.get("italic"):
        run.italic = False
    if base_format.get("underline"):
        run.underline = False
    if base_format.get("highlight"):
        set_highlight_none(run)
    font_size = base_format.get("font_size")
    if font_size:
        run.font.size = font_size


FONT_TOKEN_RE = re.compile(r"\d[\d,.\-/%:]*|[\u3400-\u9fff]+|[^\d\u3400-\u9fff]+")


def split_font_tokens(text: str) -> list[str]:
    return [match.group(0) for match in FONT_TOKEN_RE.finditer(text or "")]


def set_run_fonts(run, token: str, english_font: str, chinese_font: str) -> None:
    ascii_font = DIGIT_FONT if token and token[0].isdigit() else valid_english_font(english_font)
    chinese_font = valid_chinese_font(chinese_font)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), chinese_font)


def add_run_with_base_format(paragraph, text: str, base_format: dict[str, object], english_font: str, chinese_font: str):
    run = paragraph.add_run(text)
    apply_base_format(run, base_format)
    set_run_fonts(run, text, english_font, chinese_font)
    return run


def apply_style(run, span: FormatSpan, base_format: dict[str, object], english_font: str, chinese_font: str) -> None:
    apply_base_format(run, base_format)
    set_run_fonts(run, run.text, english_font, chinese_font)
    if span.bold:
        run.bold = True
    if span.italic:
        run.italic = True
    if span.underline:
        run.underline = span.underline
    if span.highlight_color:
        run.font.highlight_color = span.highlight_color


def add_text_runs(paragraph, text: str, base_format: dict[str, object], english_font: str, chinese_font: str, span: FormatSpan | None = None) -> None:
    for token in split_font_tokens(text):
        run = add_run_with_base_format(paragraph, token, base_format, english_font, chinese_font)
        if span is not None:
            apply_style(run, span, base_format, english_font, chinese_font)


def rewrite_paragraph(paragraph, translation: str, intervals: list[dict], english_font: str, chinese_font: str) -> None:
    intervals = sorted(intervals, key=lambda item: (item["start"], item["end"]))
    base_format = base_format_overrides(paragraph)
    clear_paragraph_runs(paragraph)
    cursor = 0
    for item in intervals:
        start = max(0, item["start"])
        end = min(len(translation), item["end"])
        if start < cursor or end <= start:
            continue
        if cursor < start:
            add_text_runs(paragraph, translation[cursor:start], base_format, english_font, chinese_font)
        add_text_runs(paragraph, translation[start:end], base_format, english_font, chinese_font, item["span"])
        cursor = end
    if cursor < len(translation):
        add_text_runs(paragraph, translation[cursor:], base_format, english_font, chinese_font)
    if not translation:
        add_run_with_base_format(paragraph, "", base_format, english_font, chinese_font)


def normalize_final_cleanup_text(text: str) -> str:
    text = strip_markdown_emphasis(text or "")
    text = text.translate(CJK_PUNCT_TRANSLATION)
    text = re.sub(r" +([,.;:!?\]\)])", r"\1", text)
    text = re.sub(r"([\[\(]) +", r"\1", text)
    text = re.sub(r"([,;:!?])(?=[A-Za-z\"'])", r"\1 ", text)
    text = re.sub(r"(?<!\d)(\.)(?=[A-Za-z\"'])", r"\1 ", text)
    text = re.sub(r"(?<=\d), +(?=\d{3}\b)", ",", text)
    text = re.sub(r"(?<=\d)\. +(?=\d)", ".", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def collect_cjk_document_paragraphs(doc: Document) -> list[dict]:
    items = []
    seen_xml_paragraphs = set()
    for index, paragraph in enumerate(iter_document_paragraphs(doc)):
        seen_xml_paragraphs.add(id(paragraph._p))
        text = paragraph_text(paragraph)
        if cjk_count(text):
            items.append(
                {
                    "id": index,
                    "display_id": f"doc:{index}",
                    "paragraph": paragraph,
                    "xml_paragraph": None,
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "",
                }
            )
            continue
        xml_text = xml_paragraph_text(paragraph._p)
        if cjk_count(xml_text):
            items.append(
                {
                    "id": index,
                    "display_id": f"doc-xml:{index}",
                    "paragraph": None,
                    "xml_paragraph": paragraph._p,
                    "text": xml_text,
                    "style": paragraph.style.name if paragraph.style else "",
                }
            )
    next_id = len(seen_xml_paragraphs)
    for part in doc.part.package.parts:
        partname = str(part.partname)
        if not (partname.startswith("/word/header") or partname.startswith("/word/footer")):
            continue
        element = getattr(part, "element", None)
        if element is None:
            continue
        for xml_index, xml_paragraph in enumerate(element.xpath(".//w:p")):
            text = xml_paragraph_text(xml_paragraph)
            if not cjk_count(text):
                continue
            item_id = next_id
            next_id += 1
            items.append(
                {
                    "id": item_id,
                    "display_id": f"{partname}:{xml_index}",
                    "paragraph": None,
                    "xml_paragraph": xml_paragraph,
                    "text": text,
                    "style": partname,
                }
            )
    return items


def build_final_doc_cleanup_batches(items: list[dict], max_chars: int = 12000, max_items: int = 18) -> list[list[dict]]:
    batches = []
    current = []
    current_chars = 0
    for item in items:
        size = len(item.get("text", "")) + 300
        if current and (current_chars + size > max_chars or len(current) >= max_items):
            batches.append(current)
            current = []
            current_chars = 0
        current.append(item)
        current_chars += size
    if current:
        batches.append(current)
    return batches


def final_doc_cleanup_batch(client: OpenAI, provider: str, model: str, memory: str, batch: list[dict], font_instruction: str) -> dict[int, str]:
    payload = [
        {
            "id": int(item["id"]),
            "style": item.get("style", ""),
            "text": item.get("text", ""),
        }
        for item in batch
    ]
    source_text = "\n".join(item["text"] for item in payload)
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Final visible Word sweep for an already translated English legal document.\n"
                "Some paragraphs still contain Chinese/CJK characters. Translate or romanize only the remaining Chinese/CJK parts while preserving the existing English text.\n"
                "Return JSON exactly like {\"items\":[{\"id\":123,\"translation\":\"...\"}]}.\n\n"
                "Rules:\n"
                "1. Return one cleaned paragraph for every input id.\n"
                "2. The returned translation must contain no Chinese/CJK characters.\n"
                "3. Do not use English Name (Chinese Name) formatting; use official English names where known, otherwise use English-only romanization or legal-English rendering.\n"
                "4. Preserve tabs, numbering, page numbers, clause references, dates, amounts, punctuation intent, and line order as much as possible.\n"
                "5. If a paragraph is a table-of-contents line, translate the title and keep the page number.\n"
                "6. Preserve bracket symbols 【】 and blanks; if Chinese appears inside 【】, translate the bracket content but keep the symbols.\n\n"
                f"FONT INSTRUCTIONS:\n{font_instruction}\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(source_text)}\n\n"
                f"TRANSLATION MEMORY:\n{memory or '(none)'}\n\n"
                f"PARAGRAPHS JSON:\n{json.dumps({'items': payload}, ensure_ascii=False)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages, retries=2)
    cleaned = {}
    for item in data.get("items", []) or []:
        try:
            item_id = int(item.get("id"))
        except (TypeError, ValueError):
            continue
        translation = str(item.get("translation", "")).strip()
        if translation:
            cleaned[item_id] = translation
    return cleaned


def cleanup_remaining_cjk_in_document(
    doc: Document,
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    log,
    progress,
    english_font: str,
    chinese_font: str,
    font_instruction: str,
    max_rounds: int = 2,
) -> tuple[set[int], list[dict]]:
    cleaned_ids: set[int] = set()
    remaining: list[dict] = []
    for round_index in range(1, max_rounds + 1):
        items = collect_cjk_document_paragraphs(doc)
        remaining = items
        if not items:
            log(f"{APP_VERSION} final Word sweep passed: no Chinese/CJK remains in visible document text.")
            return cleaned_ids, []
        log(f"{APP_VERSION} final Word sweep round {round_index}: found {len(items)} visible paragraphs with Chinese/CJK.")
        changed = False
        batches = build_final_doc_cleanup_batches(items)
        for batch_index, batch in enumerate(batches, start=1):
            progress(len(cleaned_ids), max(1, len(items)), f"{APP_VERSION} final Word Chinese cleanup {batch_index}/{len(batches)}")
            try:
                repaired = final_doc_cleanup_batch(client, provider, model, memory, batch, font_instruction)
            except TranslationCancelled:
                raise
            except Exception as exc:
                log(f"{APP_VERSION} final Word cleanup batch failed and was skipped: {exc}")
                repaired = {}
            for item in batch:
                item_id = int(item["id"])
                old_text = item.get("text", "")
                new_text = normalize_final_cleanup_text(repaired.get(item_id, ""))
                if not new_text or new_text == old_text:
                    continue
                if cjk_count(new_text) and cjk_count(new_text) >= cjk_count(old_text):
                    continue
                if item.get("paragraph") is not None:
                    rewrite_paragraph(item["paragraph"], new_text, [], english_font, chinese_font)
                elif item.get("xml_paragraph") is not None:
                    rewrite_xml_paragraph_text(item["xml_paragraph"], new_text)
                else:
                    continue
                cleaned_ids.add(item_id)
                changed = True
        if not changed:
            break
    remaining = collect_cjk_document_paragraphs(doc)
    if remaining:
        samples = ", ".join(f"{item['id']}: {compact_text(item['text'], 60)}" for item in remaining[:10])
        log(f"{APP_VERSION} final Word sweep still found Chinese/CJK after cleanup: {samples}{'...' if len(remaining) > 10 else ''}")
    return cleaned_ids, remaining


def safe_save_document(doc: Document, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f"{output_path.stem}.tmp.{int(time.time() * 1000)}{output_path.suffix}")
    doc.save(str(tmp_path))
    try:
        os.replace(str(tmp_path), str(output_path))
        return output_path
    except OSError as exc:
        if getattr(exc, "winerror", None) != 5 and getattr(exc, "errno", None) != 13:
            tmp_path.unlink(missing_ok=True)
            raise
        fallback = output_path.with_name(f"{output_path.stem}_自动另存_{time.strftime('%Y%m%d_%H%M%S')}{output_path.suffix}")
        os.replace(str(tmp_path), str(fallback))
        return fallback


def write_process_files_note(process_dir: Path) -> Path:
    note_path = process_dir / "说明_这些过程文件可删除.txt"
    note_path.write_text(
        (
            "这里保存的是翻译过程文件，用于排查问题或回看模型处理记录。\n\n"
            "最终交付给客户/同事通常只需要外层输出目录里的 Word 文件：*_复合方法英文翻译.docx。\n"
            "如果译文 Word 已经确认没有问题，本文件夹里的 .md、.json 和本说明文件都可以删除。\n"
        ),
        encoding="utf-8",
    )
    return note_path


def checklist_markdown(rows: list[dict]) -> str:
    lines = [
        "# 复合方法 Checklist",
        "",
        "| 状态 | Block | 原格式文本 | 格式 | 英文映射 | 置信度 | 备注 |",
        "|---|---:|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append(
            "| {status} | {block_id} | {source_text} | {style} | {target_text} | {confidence} | {note} |".format(
                status=row.get("status", ""),
                block_id=row.get("block_id", ""),
                source_text=str(row.get("source_text", "")).replace("|", "\\|"),
                style=str(row.get("style", "")).replace("|", "\\|"),
                target_text=str(row.get("target_text", "")).replace("|", "\\|"),
                confidence=str(row.get("confidence", "")).replace("|", "\\|"),
                note=str(row.get("note", "")).replace("|", "\\|"),
            )
        )
    return "\n".join(lines) + "\n"


def translate_docx_hybrid(
    input_path: Path,
    output_dir: Path,
    provider: str,
    api_key: str,
    base_url: str,
    model: str,
    english_font: str,
    chinese_font: str,
    log,
    translation_progress,
    review_progress,
    cancel_event: threading.Event | None = None,
) -> tuple[Path, Path, Path, Path, Path, bool]:
    set_active_cancel_event(cancel_event)
    output_dir.mkdir(parents=True, exist_ok=True)
    english_font = valid_english_font(english_font)
    chinese_font = valid_chinese_font(chinese_font)
    font_instruction = build_font_instruction(english_font, chinese_font)
    doc = Document(str(input_path))
    paragraphs = list(iter_document_paragraphs(doc))
    blocks = []
    spans_by_id: dict[str, FormatSpan] = {}

    for block_id, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        if not has_cjk(text):
            continue
        spans = extract_format_spans(paragraph, block_id)
        public_spans = []
        for span in spans:
            spans_by_id[span.span_id] = span
            public_spans.append({"span_id": span.span_id, "source_text": span.source_text, "style": span.style_label})
        block = {"id": block_id, "text": text, "meta": paragraph_meta(paragraph), "format_spans": public_spans}
        block["markdown"] = block_to_markdown(block)
        blocks.append(block)

    if not blocks:
        raise ValueError("没有在文档中发现中文内容。")

    source_markdown = "\n\n".join(block["markdown"] for block in blocks)
    client = build_client(api_key, base_url)
    report_progress(review_progress, 0, "等待翻译完成后开始复核检查")
    if ENABLE_COMPANY_NAME_GLOSSARY:
        report_progress(translation_progress, 2, "正在抽取公司/机构名表...")
        company_entries = make_company_name_glossary(client, provider, model, source_markdown, log)
        company_glossary = company_glossary_to_prompt(company_entries)
    else:
        company_entries = []
        company_glossary = "(disabled; do not force English Name (Chinese Name) formatting)"
        log(f"{APP_VERSION} 公司名强制“英文（中文）”纠错已关闭，避免复核阶段反复重翻。")
    report_progress(translation_progress, 4, "正在生成术语记忆和翻译规则...")
    memory = make_translation_memory(client, provider, model, source_markdown, log, font_instruction, company_glossary)
    chunks = build_chunks(blocks)
    log(f"{APP_VERSION} 复合方法：{len(blocks)} 个中文 block，分为 {len(chunks)} 个 Markdown 大段批次。")
    log(f"字体设置：英文 {english_font}；数字 {DIGIT_FONT}；中文 {chinese_font}。")

    translated_blocks: dict[int, str] = {}
    translated_markdown_parts = []
    all_mappings = []
    auto_repaired_blocks: set[int] = set()
    final_audit_repaired_blocks: set[int] = set()
    completed = 0
    cancelled = False
    for chunk_index, chunk in enumerate(chunks, start=1):
        if cancel_event and cancel_event.is_set():
            cancelled = True
            log("收到中止请求，停止进入下一个 Markdown 批次，开始导出当前进度。")
            break
        chunk_markdown = "\n\n".join(block["markdown"] for block in chunk)
        before = "\n\n".join(block["markdown"] for block in blocks[max(0, blocks.index(chunk[0]) - 3) : blocks.index(chunk[0])])
        after_start = blocks.index(chunk[-1]) + 1
        after = "\n\n".join(block["markdown"] for block in blocks[after_start : min(len(blocks), after_start + 3)])
        report_phase(translation_progress, 5, 99, completed, len(blocks), f"翻译 Markdown 大段 {chunk_index}/{len(chunks)}")
        try:
            translated_markdown = translate_markdown_chunk(client, provider, model, memory, chunk_markdown, before, after, font_instruction, company_glossary)
        except TranslationCancelled:
            cancelled = True
            log("收到中止请求；当前 API 批次不再等待，直接导出最近已完成进度。")
            report_progress(review_progress, 85, "正在导出最近已完成进度")
            break
        translated_markdown_parts.append(translated_markdown)
        chunk_translations = parse_translated_markdown(translated_markdown)
        for block in chunk:
            block_id = int(block["id"])
            if block_id in chunk_translations:
                chunk_translations[block_id] = normalize_translation_against_source(block["text"], chunk_translations[block_id])
        translated_blocks.update(chunk_translations)
        completed += len(chunk)
        if cancel_event and cancel_event.is_set():
            cancelled = True
            log("当前 Markdown 批次已完成；根据中止请求，开始导出当前进度。")
            report_progress(review_progress, 85, "已中止，正在导出当前进度")
            break
        report_phase(translation_progress, 5, 99, completed, len(blocks), f"大段 {chunk_index}/{len(chunks)} 已完成")
    if not cancelled:
        report_complete(translation_progress, "翻译完成，开始复核检查")

    if not cancelled:
        try:
            auto_repaired_blocks.update(
                repair_all_translations(
                    client,
                    provider,
                    model,
                    memory,
                    blocks,
                    translated_blocks,
                    log,
                    make_phase_progress(review_progress, 0, 35),
                    font_instruction,
                    company_glossary,
                )
            )
        except TranslationCancelled:
            cancelled = True
            log("收到中止请求；漏翻自修复当前批次不再等待，直接导出最近已完成进度。")
            report_progress(review_progress, 85, "正在导出最近已完成进度")
    if not cancelled:
        try:
            final_audit_repaired_blocks.update(
                final_llm_audit_and_repair_translations(
                    client,
                    provider,
                    model,
                    memory,
                    blocks,
                    translated_blocks,
                    log,
                    make_phase_progress(review_progress, 35, 65),
                    font_instruction,
                    company_entries,
                    company_glossary,
                )
            )
        except TranslationCancelled:
            cancelled = True
            log("收到中止请求；最终复核当前批次不再等待，直接导出最近已完成进度。")
            report_progress(review_progress, 85, "正在导出最近已完成进度")
    all_mappings = []
    for chunk_index, chunk in enumerate(chunks, start=1):
        ready_chunk = [block for block in chunk if translated_blocks.get(int(block["id"]), "").strip()]
        if not ready_chunk:
            continue
        report_phase(review_progress, 65, 85, chunk_index - 1, len(chunks), f"映射格式 {chunk_index}/{len(chunks)}")
        chunk_translations = {int(block["id"]): translated_blocks.get(int(block["id"]), "") for block in ready_chunk}
        try:
            all_mappings.extend(map_format_spans(client, provider, model, memory, ready_chunk, chunk_translations, font_instruction))
        except TranslationCancelled:
            cancelled = True
            log("收到中止请求；格式映射当前批次不再等待，直接导出最近已完成进度。")
            break
        report_phase(review_progress, 65, 85, chunk_index, len(chunks), f"格式映射 {chunk_index}/{len(chunks)} 已完成")

    mappings_by_span = {str(mapping.get("span_id")): mapping for mapping in all_mappings if mapping.get("span_id")}
    checklist = []
    if cancelled:
        checklist.append(
            {
                "status": "CANCELLED",
                "block_id": "",
                "source_text": "",
                "style": "",
                "target_text": "",
                "confidence": "",
                "note": "用户中止；已导出当前已完成批次，未完成 block 保留原文。",
            }
        )
    for index, block in enumerate(blocks, start=1):
        if index == 1 or index == len(blocks) or index % 25 == 0:
            report_phase(review_progress, 85, 94, index - 1, len(blocks), f"正在写入译文和格式 {index}/{len(blocks)}")
        block_id = block["id"]
        translation = normalize_translation_against_source(block["text"], translated_blocks.get(block_id, "").strip())
        intervals = []
        occupied = []
        if not translation:
            checklist.append({"status": "MISSING_TRANSLATION", "block_id": block_id, "note": "模型没有返回该 block 的译文。"})
            continue
        if block_id in auto_repaired_blocks:
            checklist.append(
                {
                    "status": "AUTO_REPAIRED_TRANSLATION",
                    "block_id": block_id,
                    "source_text": "",
                    "style": "",
                    "target_text": "",
                    "confidence": "deterministic",
                    "note": "自检查发现漏翻或中文残留后，已在导出前自动重翻该 block。",
                }
            )
        if block_id in final_audit_repaired_blocks:
            checklist.append(
                {
                    "status": "FINAL_LLM_AUDIT_REPAIRED",
                    "block_id": block_id,
                    "source_text": "",
                    "style": "",
                    "target_text": "",
                    "confidence": "llm-audit",
                    "note": f"{APP_VERSION} final LLM audit found untranslated Chinese/CJK legal content and repaired this block before export.",
                }
            )
        for public_span in block["format_spans"]:
            span_id = public_span["span_id"]
            span = spans_by_id[span_id]
            mapping = mappings_by_span.get(span_id)
            row = {
                "status": "",
                "block_id": block_id,
                "source_text": span.source_text,
                "style": span.style_label,
                "target_text": "",
                "confidence": "",
                "note": "",
            }
            if not mapping:
                fallback_found = None
                fallback_target = ""
                for candidate in source_span_fallback_targets(span.source_text):
                    fallback_target = normalize_legal_english(candidate)
                    fallback_found = find_available_span(translation, fallback_target, occupied)
                    if fallback_found:
                        break
                if fallback_found:
                    start, end = fallback_found
                    intervals.append({"start": start, "end": end, "span": span})
                    occupied.append((start, end))
                    row["target_text"] = fallback_target
                    row["confidence"] = "deterministic"
                    row["status"] = "APPLIED_FALLBACK"
                    row["note"] = "Applied deterministic legal-reference mapping because the model omitted this span."
                    checklist.append(row)
                    continue
                row["status"] = "MISSING_MAPPING"
                row["note"] = "模型没有返回该格式 span 的英文映射。"
                checklist.append(row)
                continue
            target_text = normalize_legal_english(str(mapping.get("target_text", "")).strip())
            row["target_text"] = target_text
            row["confidence"] = str(mapping.get("confidence", "")).strip()
            row["note"] = str(mapping.get("note", "")).strip()
            found = find_available_span(translation, target_text, occupied)
            if not found:
                for candidate in source_span_fallback_targets(span.source_text):
                    target_text = normalize_legal_english(candidate)
                    found = find_available_span(translation, target_text, occupied)
                    if found:
                        row["target_text"] = target_text
                        row["confidence"] = "deterministic"
                        row["note"] = (row["note"] + " " if row["note"] else "") + "Used deterministic legal-reference fallback."
                        break
            if not found:
                row["status"] = "TARGET_NOT_FOUND"
                row["note"] = (row["note"] + " " if row["note"] else "") + "英文映射未在译文中精确找到。"
                checklist.append(row)
                continue
            start, end = found
            intervals.append({"start": start, "end": end, "span": span})
            occupied.append((start, end))
            row["status"] = "APPLIED"
            checklist.append(row)
        if needs_translation_repair(block["text"], translation):
            checklist.append({"status": "HAS_CHINESE_IN_TRANSLATION", "block_id": block_id, "note": "译文仍疑似存在漏翻或大段中文残留。"})
        rewrite_paragraph(paragraphs[block_id], translation, intervals, english_font, chinese_font)
    if not cancelled:
        try:
            cleaned_doc_paragraphs, remaining_doc_paragraphs = cleanup_remaining_cjk_in_document(
                doc,
                client,
                provider,
                model,
                memory,
                log,
                make_phase_progress(review_progress, 94, 98),
                english_font,
                chinese_font,
                font_instruction,
            )
            for paragraph_id in sorted(cleaned_doc_paragraphs):
                checklist.append(
                    {
                        "status": "FINAL_DOC_CJK_CLEANUP",
                        "block_id": f"doc:{paragraph_id}",
                        "source_text": "",
                        "style": "",
                        "target_text": "",
                        "confidence": "llm-final-sweep",
                        "note": "最终 Word 可见文本扫尾发现中文/CJK，并已在导出前重翻清理。",
                    }
                )
            for item in remaining_doc_paragraphs:
                checklist.append(
                    {
                        "status": "HAS_CHINESE_IN_FINAL_DOC",
                        "block_id": item.get("display_id", f"doc:{item['id']}"),
                        "source_text": compact_text(item.get("text", ""), 120),
                        "style": item.get("style", ""),
                        "target_text": "",
                        "confidence": "",
                        "note": "最终 Word 可见文本扫尾后仍检测到中文/CJK，请人工复核该段。",
                    }
                )
        except TranslationCancelled:
            cancelled = True
            log("收到中止请求；最终 Word 中文扫尾当前批次不再等待，直接导出最近已完成进度。")
            report_progress(review_progress, 98, "正在导出最近已完成进度")
        except Exception as exc:
            log(f"{APP_VERSION} final Word Chinese cleanup failed; exporting current document and recording the issue. Reason: {exc}")
            for item in collect_cjk_document_paragraphs(doc):
                checklist.append(
                    {
                        "status": "HAS_CHINESE_IN_FINAL_DOC",
                        "block_id": item.get("display_id", f"doc:{item['id']}"),
                        "source_text": compact_text(item.get("text", ""), 120),
                        "style": item.get("style", ""),
                        "target_text": "",
                        "confidence": "",
                        "note": "最终 Word 可见文本扫尾执行失败，该段仍检测到中文/CJK。",
                    }
                )
    report_progress(review_progress, 98, "正在导出 Word 和明细文件...")

    base = input_path.stem
    suffix = "_已中止" if cancelled else ""
    process_dir = output_dir / "过程文件"
    process_dir.mkdir(parents=True, exist_ok=True)
    docx_path = safe_save_document(doc, output_dir / f"{base}_复合方法英文翻译.docx")
    source_md_path = process_dir / f"{base}_source_blocks.md"
    translated_md_path = process_dir / f"{base}_translated_blocks.md"
    checklist_path = process_dir / f"{base}_复合方法checklist.md"
    json_path = process_dir / f"{base}_复合方法明细.json"
    if cancelled:
        cancelled_docx_path = docx_path.with_name(f"{docx_path.stem}{suffix}{docx_path.suffix}")
        os.replace(str(docx_path), str(cancelled_docx_path))
        docx_path = cancelled_docx_path
        source_md_path = source_md_path.with_name(f"{source_md_path.stem}{suffix}{source_md_path.suffix}")
        translated_md_path = translated_md_path.with_name(f"{translated_md_path.stem}{suffix}{translated_md_path.suffix}")
        checklist_path = checklist_path.with_name(f"{checklist_path.stem}{suffix}{checklist_path.suffix}")
        json_path = json_path.with_name(f"{json_path.stem}{suffix}{json_path.suffix}")

    write_process_files_note(process_dir)
    source_md_path.write_text(source_markdown, encoding="utf-8")
    translated_md_path.write_text(build_translated_blocks_markdown(blocks, translated_blocks), encoding="utf-8")
    checklist_path.write_text(checklist_markdown(checklist), encoding="utf-8")
    json_path.write_text(
        json.dumps(
            {
                "cancelled": cancelled,
                "company_glossary": company_entries,
                "blocks": blocks,
                "mappings": all_mappings,
                "checklist": checklist,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    report_complete(review_progress, "已中止并导出当前进度" if cancelled else "全部完成")
    return docx_path, source_md_path, translated_md_path, checklist_path, json_path, cancelled


class HybridApp:
    def __init__(self, root):
        self.root = root
        self.main_thread_id = threading.get_ident()
        self.root.title("复合方法 - Markdown 结构 + run 格式映射")
        self.root.geometry("880x700")
        self.root.minsize(800, 640)

        config = load_config()
        provider = config.get("provider", "DeepSeek")
        if provider not in PROVIDER_DEFAULTS:
            provider = "DeepSeek"
        defaults = PROVIDER_DEFAULTS[provider]
        self.provider = StringVar(value=provider)
        self.api_key = StringVar(value=config.get("api_key", ""))
        self.base_url = StringVar(value=config.get("base_url", defaults["base_url"]))
        self.model = StringVar(value=config.get("model", defaults["model"]))
        self.english_font = StringVar(value=valid_english_font(config.get("english_font", DEFAULT_ENGLISH_FONT)))
        self.chinese_font = StringVar(value=valid_chinese_font(config.get("chinese_font", DEFAULT_CHINESE_FONT)))
        self.file_path = StringVar(value="")
        self.output_dir = StringVar(value=str(OUTPUT_DIR))
        self.remember_key = BooleanVar(value=bool(config.get("api_key")))
        self.cancel_event = threading.Event()

        self.canvas = Canvas(root, highlightthickness=0)
        self.page_scrollbar = Scrollbar(root, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.page_scrollbar.set)
        self.page_scrollbar.pack(side=RIGHT, fill="y")
        self.canvas.pack(side=LEFT, fill=BOTH, expand=True)
        self.content = Frame(self.canvas)
        self.content_window = self.canvas.create_window((0, 0), window=self.content, anchor="nw")
        self.content.bind("<Configure>", self.on_content_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)
        self.canvas.bind_all("<MouseWheel>", self.on_mousewheel)

        top = Frame(self.content, padx=16, pady=12)
        top.pack(fill=X)
        Label(top, text="API 类型").pack(anchor="w")
        OptionMenu(top, self.provider, *PROVIDER_DEFAULTS.keys(), command=self.on_provider_change).pack(fill=X, pady=(4, 8))
        self.key_label = Label(top, text=defaults["key_label"])
        self.key_label.pack(anchor="w")
        Entry(top, textvariable=self.api_key, show="*", width=90).pack(fill=X, pady=(4, 8))
        Label(top, text="Base URL").pack(anchor="w")
        Entry(top, textvariable=self.base_url, width=90).pack(fill=X, pady=(4, 8))
        Label(top, text="模型").pack(anchor="w")
        Entry(top, textvariable=self.model, width=90).pack(fill=X, pady=(4, 8))
        font_row = Frame(top)
        font_row.pack(fill=X, pady=(0, 8))
        left_font = Frame(font_row)
        left_font.pack(side=LEFT, fill=X, expand=True)
        Label(left_font, text="英文字体（字母）").pack(anchor="w")
        OptionMenu(left_font, self.english_font, *ENGLISH_FONT_OPTIONS).pack(fill=X, pady=(4, 0))
        right_font = Frame(font_row)
        right_font.pack(side=RIGHT, fill=X, expand=True, padx=(12, 0))
        Label(right_font, text=f"中文字体（数字固定 {DIGIT_FONT}）").pack(anchor="w")
        OptionMenu(right_font, self.chinese_font, *CHINESE_FONT_OPTIONS).pack(fill=X, pady=(4, 0))
        Checkbutton(top, text="记住 API Key（明文保存在本机复合方法目录）", variable=self.remember_key).pack(anchor="w")

        file_frame = Frame(self.content, padx=16, pady=8)
        file_frame.pack(fill=X)
        Label(file_frame, text="合同文件（.docx）").pack(anchor="w")
        row = Frame(file_frame)
        row.pack(fill=X, pady=(4, 8))
        Entry(row, textvariable=self.file_path).pack(side=LEFT, fill=X, expand=True)
        Button(row, text="选择文件", command=self.choose_file, width=12).pack(side=RIGHT, padx=(8, 0))
        Label(file_frame, text="输出目录").pack(anchor="w")
        out_row = Frame(file_frame)
        out_row.pack(fill=X, pady=(4, 0))
        Entry(out_row, textvariable=self.output_dir).pack(side=LEFT, fill=X, expand=True)
        Button(out_row, text="选择目录", command=self.choose_output_dir, width=12).pack(side=RIGHT, padx=(8, 0))

        self.drop_label = Label(self.content, text="把 Word 合同拖到这里\n输出：DOCX + source Markdown + translated Markdown + checklist + JSON", relief="groove", height=4)
        self.drop_label.pack(fill=X, padx=16, pady=8)
        if DND_AVAILABLE:
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind("<<Drop>>", self.on_drop)

        progress_frame = Frame(self.content, padx=16, pady=8)
        progress_frame.pack(fill=X)
        self.translation_progress_text = StringVar(value="翻译进度：0/100（0.0%）")
        self.translation_current_text = StringVar(value="翻译：尚未开始")
        self.review_progress_text = StringVar(value="复核检查：0/100（0.0%）")
        self.review_current_text = StringVar(value="复核：等待翻译完成")
        Label(progress_frame, textvariable=self.translation_progress_text).pack(anchor="w")
        self.translation_bar = Progressbar(progress_frame, maximum=100)
        self.translation_bar.pack(fill=X, pady=(4, 4))
        Label(progress_frame, textvariable=self.translation_current_text, wraplength=820, justify=LEFT).pack(anchor="w")
        Label(progress_frame, textvariable=self.review_progress_text).pack(anchor="w", pady=(8, 0))
        self.review_bar = Progressbar(progress_frame, maximum=100)
        self.review_bar.pack(fill=X, pady=(4, 4))
        Label(progress_frame, textvariable=self.review_current_text, wraplength=820, justify=LEFT).pack(anchor="w")

        action = Frame(self.content, padx=16, pady=8)
        action.pack(fill=X)
        self.start_button = Button(action, text="开始复合方法翻译", command=self.start, height=2)
        self.start_button.pack(side=LEFT, fill=X, expand=True)
        self.stop_button = Button(action, text="中止并导出当前进度", command=self.request_cancel, height=2, state="disabled")
        self.stop_button.pack(side=RIGHT, padx=(8, 0))

        self.log_box = Text(self.content, height=14, wrap="word")
        self.log_box.pack(fill=BOTH, expand=True, padx=16, pady=(4, 16))
        self.log(f"{APP_VERSION} 复合方法：Markdown 负责结构和上下文，run 结构负责格式抽取，模型负责格式片段到英文短语的映射。")

    def on_content_configure(self, _event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_canvas_configure(self, event):
        self.canvas.itemconfigure(self.content_window, width=event.width)

    def on_mousewheel(self, event):
        if event.delta:
            self.canvas.yview_scroll(int(-event.delta / 120), "units")

    def on_provider_change(self, provider):
        defaults = PROVIDER_DEFAULTS[provider]
        self.key_label.config(text=defaults["key_label"])
        self.base_url.set(defaults["base_url"])
        self.model.set(defaults["model"])

    def choose_file(self):
        path = filedialog.askopenfilename(title="选择 Word 合同", filetypes=[("Word 文档", "*.docx")])
        if path:
            self.file_path.set(path)

    def choose_output_dir(self):
        path = filedialog.askdirectory(title="选择输出目录", initialdir=self.output_dir.get() or str(OUTPUT_DIR))
        if path:
            self.output_dir.set(path)

    def on_drop(self, event):
        paths = self.root.tk.splitlist(event.data)
        if paths:
            self.file_path.set(paths[0])

    def log(self, text: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.log, text)
            return
        self.log_box.insert(END, text + "\n")
        self.log_box.see(END)
        self.root.update_idletasks()

    def update_translation_progress(self, done: int, total: int, current: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.update_translation_progress, done, total, current)
            return
        percent = 0 if total <= 0 else max(0, min(100, done * 100 / total))
        self.translation_bar["value"] = percent
        self.translation_progress_text.set(f"翻译进度：{done}/{total}（{percent:.1f}%）")
        self.translation_current_text.set(f"翻译：{current}")
        self.root.update_idletasks()

    def update_review_progress(self, done: int, total: int, current: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.update_review_progress, done, total, current)
            return
        percent = 0 if total <= 0 else max(0, min(100, done * 100 / total))
        self.review_bar["value"] = percent
        self.review_progress_text.set(f"复核检查：{done}/{total}（{percent:.1f}%）")
        self.review_current_text.set(f"复核：{current}")
        self.root.update_idletasks()

    def start(self):
        provider = self.provider.get().strip()
        api_key = self.api_key.get().strip()
        base_url = self.base_url.get().strip()
        model = self.model.get().strip()
        english_font = valid_english_font(self.english_font.get().strip())
        chinese_font = valid_chinese_font(self.chinese_font.get().strip())
        input_path = Path(self.file_path.get().strip().strip('"'))
        output_dir = Path(self.output_dir.get().strip().strip('"'))
        if not api_key:
            messagebox.showerror("缺少 API Key", f"请先输入 {provider} API Key。")
            return
        if not model:
            messagebox.showerror("缺少模型", "请填写模型名称。")
            return
        if not input_path.exists() or input_path.suffix.lower() != ".docx":
            messagebox.showerror("文件不支持", "请选择存在的 .docx 文件。")
            return
        if provider == "DeepSeek" and not base_url:
            base_url = PROVIDER_DEFAULTS["DeepSeek"]["base_url"]
            self.base_url.set(base_url)
        self.english_font.set(english_font)
        self.chinese_font.set(chinese_font)
        save_config(provider, api_key if self.remember_key.get() else "", base_url, model, english_font, chinese_font)
        self.cancel_event.clear()
        self.start_button.config(state="disabled")
        self.stop_button.config(state="normal")
        self.update_translation_progress(0, PROGRESS_TOTAL, "准备开始")
        self.update_review_progress(0, PROGRESS_TOTAL, "等待翻译完成")
        thread = threading.Thread(
            target=self.worker,
            args=(input_path, output_dir, provider, api_key, base_url, model, english_font, chinese_font),
            daemon=True,
        )
        thread.start()

    def request_cancel(self):
        self.cancel_event.set()
        self.stop_button.config(state="disabled")
        self.log("已请求中止；将不再等待当前 API 批次返回，直接导出最近已完成进度。")
        self.update_review_progress(85, PROGRESS_TOTAL, "正在中止并导出最近已完成进度...")

    def worker(
        self,
        input_path: Path,
        output_dir: Path,
        provider: str,
        api_key: str,
        base_url: str,
        model: str,
        english_font: str,
        chinese_font: str,
    ):
        try:
            paths = translate_docx_hybrid(
                input_path,
                output_dir,
                provider,
                api_key,
                base_url,
                model,
                english_font,
                chinese_font,
                self.log,
                self.update_translation_progress,
                self.update_review_progress,
                self.cancel_event,
            )
            self.root.after(0, self.finish_success, paths[:-1], paths[-1])
        except TranslationCancelled:
            self.log("已中止；当前还没有完成可导出的翻译批次。")
            self.root.after(0, self.finish_error, "已中止；当前还没有完成可导出的翻译批次。")
        except Exception as exc:
            self.log("发生错误：")
            self.log(str(exc))
            self.log(traceback.format_exc())
            self.root.after(0, self.finish_error, str(exc))

    def finish_success(self, paths, cancelled: bool):
        self.start_button.config(state="normal")
        self.stop_button.config(state="disabled")
        title = "已中止并导出" if cancelled else "完成"
        messagebox.showinfo(
            title,
            "已输出：\n"
            + "\n".join(str(path) for path in paths)
            + "\n\n说明：输出目录里的“过程文件”文件夹只用于排查问题；如果 Word 译文确认没问题，可以删除。",
        )

    def finish_error(self, error: str):
        self.start_button.config(state="normal")
        self.stop_button.config(state="disabled")
        messagebox.showerror("翻译失败", error)


def main():
    root_class = TkinterDnD.Tk if DND_AVAILABLE else Tk
    root = root_class()
    HybridApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
