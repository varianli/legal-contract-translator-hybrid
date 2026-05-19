import json
import html
import os
import re
import threading
import time
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from tkinter import BOTH, END, LEFT, RIGHT, X, BooleanVar, Button, Checkbutton, DoubleVar, Entry, Frame, Label, OptionMenu, StringVar, Text, Tk, filedialog, messagebox
from tkinter.ttk import Progressbar

from docx import Document
from openai import OpenAI

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD

    DND_AVAILABLE = True
except Exception:
    TkinterDnD = None
    DND_FILES = None
    DND_AVAILABLE = False


APP_DIR = Path(__file__).resolve().parent
CONFIG_PATH = APP_DIR / ".translator_config.json"
MODE_RUN_CONCURRENT = "并发 run 保格式模式（推荐）"
MODE_FAST = "快速批量标签模式（不推荐大文档）"
MODE_CAREFUL = "逐段上下文模式（更稳）"
PARALLEL_WORKERS = 8
PROVIDER_DEFAULTS = {
    "OpenAI": {
        "base_url": "",
        "model": "gpt-4.1",
        "key_label": "OpenAI API Key",
    },
    "DeepSeek": {
        "base_url": "https://api.deepseek.com",
        "model": "deepseek-v4-flash",
        "key_label": "DeepSeek API Key",
    },
}


SYSTEM_PROMPT = """You are a senior bilingual legal translator.
Translate Chinese legal contracts into precise, formal legal English.
Preserve legal meaning, defined terms, numbering, article references, dates, names, amounts, punctuation intent, and placeholders.
Do not summarize, omit, add explanations, or provide commentary.
Return only valid JSON when JSON is requested."""


def has_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text or ""))


def compact_text(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if len(text) <= limit:
        return text
    half = max(1, limit // 2)
    return text[:half] + "\n...[middle omitted for context window]...\n" + text[-half:]


def load_config() -> dict:
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def save_config(provider: str, api_key: str, base_url: str, model: str, translation_mode: str) -> None:
    CONFIG_PATH.write_text(
        json.dumps(
            {
                "provider": provider,
                "api_key": api_key,
                "base_url": base_url,
                "model": model,
                "translation_mode": translation_mode,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def iter_table_paragraphs(table, seen):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_cell_paragraphs(cell, seen)


def iter_cell_paragraphs(cell, seen):
    for paragraph in cell.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in cell.tables:
        yield from iter_table_paragraphs(table, seen)


def iter_part_paragraphs(part, seen):
    for paragraph in part.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in part.tables:
        yield from iter_table_paragraphs(table, seen)


def iter_document_paragraphs(doc: Document):
    seen = set()
    yield from iter_part_paragraphs(doc, seen)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            yield from iter_part_paragraphs(part, seen)


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def edge_space_preserving(original: str, translated: str) -> str:
    if not original:
        return translated
    if original.strip() == "":
        return original
    leading = re.match(r"^\s*", original).group(0)
    trailing = re.search(r"\s*$", original).group(0)
    text = translated or ""
    if leading and not text.startswith(leading):
        text = leading + text.lstrip()
    if trailing and not text.endswith(trailing):
        text = text.rstrip() + trailing
    return text


def parse_json_object(raw: str) -> dict:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r"```$", "", raw).strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("The model did not return a JSON object.")
    return json.loads(raw[start : end + 1])


def build_client(api_key: str, base_url: str) -> OpenAI:
    kwargs = {"api_key": api_key, "timeout": 90}
    if base_url.strip():
        kwargs["base_url"] = base_url.strip()
    return OpenAI(**kwargs)


def completion_kwargs(provider: str, model: str, messages: list[dict], use_response_format: bool) -> dict:
    kwargs = {
        "model": model,
        "messages": messages,
        "temperature": 0.1,
    }
    if use_response_format:
        kwargs["response_format"] = {"type": "json_object"}
    return kwargs


def chat_json(client: OpenAI, provider: str, model: str, messages: list[dict], retries: int = 3) -> dict:
    last_error = None
    for attempt in range(retries):
        use_response_format_options = [True, False] if provider == "DeepSeek" else [True]
        for use_response_format in use_response_format_options:
            try:
                response = client.chat.completions.create(
                    **completion_kwargs(provider, model, messages, use_response_format)
                )
                return parse_json_object(response.choices[0].message.content)
            except Exception as exc:
                last_error = exc
        if attempt + 1 < retries:
            time.sleep(2 + attempt * 3)
    raise last_error


def make_document_memory(client: OpenAI, provider: str, model: str, all_text: list[str], log) -> str:
    source = "\n".join(t for t in all_text if has_cjk(t))
    if not source.strip():
        return ""
    source = compact_text(source, 55000)
    log("正在读取整份合同，提取术语表和上下文...")
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Create a concise bilingual legal translation memory for the contract below.\n"
                "Return JSON with one key named memory. The value should include: parties, defined terms, recurring legal phrases, "
                "capitalized term rules, and translation preferences. Keep it under 1200 Chinese characters plus English terms.\n\n"
                f"CONTRACT TEXT:\n{source}"
            ),
        },
    ]
    try:
        data = chat_json(client, provider, model, messages, retries=2)
        return str(data.get("memory", "")).strip()
    except Exception as exc:
        log(f"术语表提取失败，将继续逐段翻译。原因：{exc}")
        return ""


def translate_paragraph(client: OpenAI, provider: str, model: str, memory: str, previous_text: str, current_text: str, next_text: str, paragraph) -> dict[int, str]:
    runs = []
    for idx, run in enumerate(paragraph.runs):
        text = run.text
        if text:
            runs.append({"id": idx, "text": text})
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate the CURRENT paragraph from Chinese to formal legal English while preserving DOCX run boundaries.\n"
                "Rules:\n"
                "1. Return JSON exactly like {\"runs\":[{\"id\":0,\"text\":\"...\"}]}.\n"
                "2. Return the same run ids you receive. Do not merge, split, remove, or invent ids.\n"
                "3. Translate each run using the whole paragraph and neighboring paragraphs as context, so the concatenated runs read naturally.\n"
                "4. Keep whitespace-only runs unchanged. Keep already-English names, citations, section numbers, dates, and amounts unless they need standard legal-English formatting.\n"
                "5. Translate all Chinese text, including bracketed placeholders and draft notes. Do not leave Chinese characters in the output.\n"
                "6. Do not output comments, Markdown, or explanations.\n\n"
                f"DOCUMENT MEMORY:\n{memory or '(none)'}\n\n"
                f"PREVIOUS PARAGRAPH:\n{compact_text(previous_text, 1500)}\n\n"
                f"CURRENT PARAGRAPH FULL TEXT:\n{compact_text(current_text, 3000)}\n\n"
                f"NEXT PARAGRAPH:\n{compact_text(next_text, 1500)}\n\n"
                f"CURRENT RUNS JSON:\n{json.dumps(runs, ensure_ascii=False)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages)
    translated = {}
    for item in data.get("runs", []):
        try:
            run_id = int(item["id"])
            translated[run_id] = str(item.get("text", ""))
        except Exception:
            continue
    return translated


def translate_run_text(client: OpenAI, provider: str, model: str, memory: str, previous_text: str, current_text: str, next_text: str, run_text: str, run_index: int) -> str:
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate this single DOCX formatting run from Chinese to formal legal English.\n"
                "Return JSON exactly like {\"text\":\"...\"}.\n"
                "The run may be a sentence fragment because Word uses runs to preserve bold, underline, highlight, and other formatting.\n"
                "Use the whole paragraph and neighboring paragraphs as context so the fragment fits the surrounding English.\n"
                "Do not leave Chinese characters. Do not add explanations.\n\n"
                f"DOCUMENT MEMORY:\n{memory or '(none)'}\n\n"
                f"PREVIOUS PARAGRAPH:\n{compact_text(previous_text, 1500)}\n\n"
                f"CURRENT PARAGRAPH FULL TEXT:\n{compact_text(current_text, 4000)}\n\n"
                f"RUN INDEX:\n{run_index}\n\n"
                f"RUN TEXT:\n{run_text}\n\n"
                f"NEXT PARAGRAPH:\n{compact_text(next_text, 1500)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages, retries=2)
    return str(data.get("text", "")).strip()


def translate_paragraph_runs_strict(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    previous_text: str,
    current_text: str,
    next_text: str,
    run_texts: list[str],
) -> tuple[dict[int, str], list[str]]:
    warnings = []
    run_items = [{"id": idx, "text": text} for idx, text in enumerate(run_texts) if text]
    required_ids = {item["id"] for item in run_items}
    translated = {}

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate the CURRENT paragraph from Chinese to formal legal English while preserving DOCX run boundaries.\n"
                "Return JSON exactly like {\"runs\":[{\"id\":0,\"text\":\"...\"}]}.\n"
                "Critical rules:\n"
                "1. Return every input run id exactly once. Do not merge, split, remove, or invent ids.\n"
                "2. Translate each run using the whole paragraph and neighboring paragraphs as context.\n"
                "3. The final concatenated runs must contain no Chinese characters.\n"
                "4. Keep already-English names, citations, section numbers, dates, and amounts unless standard legal-English formatting requires adjustment.\n"
                "5. Do not output comments, Markdown, or explanations.\n\n"
                f"DOCUMENT MEMORY:\n{memory or '(none)'}\n\n"
                f"PREVIOUS PARAGRAPH:\n{compact_text(previous_text, 1500)}\n\n"
                f"CURRENT PARAGRAPH FULL TEXT:\n{compact_text(current_text, 5000)}\n\n"
                f"NEXT PARAGRAPH:\n{compact_text(next_text, 1500)}\n\n"
                f"CURRENT RUNS JSON:\n{json.dumps(run_items, ensure_ascii=False)}"
            ),
        },
    ]

    try:
        data = chat_json(client, provider, model, messages, retries=2)
        for item in data.get("runs", []):
            run_id = int(item["id"])
            if run_id in required_ids:
                translated[run_id] = str(item.get("text", ""))
    except Exception as exc:
        warnings.append(f"段落 run 批量翻译失败，改为逐 run 补译：{exc}")

    missing_ids = required_ids - set(translated)
    cjk_ids = {idx for idx, text in translated.items() if idx in required_ids and has_cjk(text)}
    fallback_ids = sorted(missing_ids | cjk_ids)

    if fallback_ids:
        warnings.append(f"补译 run：{fallback_ids}")
    for run_id in fallback_ids:
        original = run_texts[run_id]
        if not has_cjk(original):
            translated[run_id] = original
            continue
        best = ""
        for _ in range(2):
            best = translate_run_text(client, provider, model, memory, previous_text, current_text, next_text, original, run_id)
            if best and not has_cjk(best):
                break
        translated[run_id] = best or original
        if has_cjk(translated[run_id]):
            warnings.append(f"run {run_id} 补译后仍含中文，请人工复核。")

    for run_id in required_ids:
        translated.setdefault(run_id, run_texts[run_id])
    return translated, warnings


def translate_paragraph_task(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    texts: list[str],
    run_texts_by_paragraph: list[list[str]],
    paragraph_index: int,
) -> tuple[int, dict[int, str], list[str]]:
    current = texts[paragraph_index]
    previous_text = texts[paragraph_index - 1] if paragraph_index > 0 else ""
    next_text = texts[paragraph_index + 1] if paragraph_index + 1 < len(texts) else ""
    translated, warnings = translate_paragraph_runs_strict(
        client,
        provider,
        model,
        memory,
        previous_text,
        current,
        next_text,
        run_texts_by_paragraph[paragraph_index],
    )
    return paragraph_index, translated, warnings


RUN_TAG_RE = re.compile(r"<r\s+id=[\"'](\d+)[\"']>(.*?)</r>", re.DOTALL)


def paragraph_to_tagged_text(paragraph) -> str:
    pieces = []
    for idx, run in enumerate(paragraph.runs):
        if run.text:
            pieces.append(f'<r id="{idx}">{html.escape(run.text, quote=False)}</r>')
    return "".join(pieces)


def tagged_text_to_runs(tagged_text: str) -> dict[int, str]:
    tagged_text = tagged_text or ""
    matches = list(RUN_TAG_RE.finditer(tagged_text))
    if not matches and "&lt;r" in tagged_text:
        tagged_text = html.unescape(tagged_text)
        matches = list(RUN_TAG_RE.finditer(tagged_text))
    translated = {}
    for match in matches:
        translated[int(match.group(1))] = html.unescape(match.group(2))
    return translated


def nonempty_run_ids(paragraph) -> set[int]:
    return {idx for idx, run in enumerate(paragraph.runs) if run.text}


def get_run_texts(paragraph) -> list[str]:
    return [run.text for run in paragraph.runs]


def restore_run_texts(paragraph, run_texts: list[str]) -> None:
    for run, text in zip(paragraph.runs, run_texts):
        run.text = text


def apply_runs_translation(paragraph, translated_runs: dict[int, str]) -> int:
    for idx, run in enumerate(paragraph.runs):
        if idx in translated_runs:
            run.text = edge_space_preserving(run.text, translated_runs[idx])
    return len(translated_runs)


def apply_tagged_translation(paragraph, tagged_text: str) -> int:
    translated_runs = tagged_text_to_runs(tagged_text)
    return apply_runs_translation(paragraph, translated_runs)


def set_plain_paragraph_translation(paragraph, text: str) -> None:
    nonempty_runs = [run for run in paragraph.runs if run.text]
    if not nonempty_runs:
        paragraph.add_run(text)
        return
    first = nonempty_runs[0]
    first.text = edge_space_preserving(first.text, text)
    for run in nonempty_runs[1:]:
        run.text = ""


def translate_plain_paragraph(client: OpenAI, provider: str, model: str, memory: str, previous_text: str, current_text: str, next_text: str) -> str:
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate the CURRENT paragraph from Chinese to precise formal legal English.\n"
                "Return JSON exactly like {\"text\":\"...\"}.\n"
                "Translate all Chinese text, including bracketed placeholders and draft notes. Do not leave Chinese characters.\n"
                "Do not summarize, omit, add explanations, or output Markdown.\n\n"
                f"DOCUMENT MEMORY:\n{memory or '(none)'}\n\n"
                f"PREVIOUS PARAGRAPH:\n{compact_text(previous_text, 1500)}\n\n"
                f"CURRENT PARAGRAPH:\n{compact_text(current_text, 5000)}\n\n"
                f"NEXT PARAGRAPH:\n{compact_text(next_text, 1500)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages, retries=2)
    return str(data.get("text", "")).strip()


def translate_paragraph_with_fallback(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    texts: list[str],
    paragraphs: list,
    paragraph_index: int,
    log,
    initial_tagged_text: str | None = None,
) -> bool:
    paragraph = paragraphs[paragraph_index]
    original_run_texts = get_run_texts(paragraph)
    required_ids = nonempty_run_ids(paragraph)
    current = texts[paragraph_index]
    previous_text = texts[paragraph_index - 1] if paragraph_index > 0 else ""
    next_text = texts[paragraph_index + 1] if paragraph_index + 1 < len(texts) else ""

    if initial_tagged_text:
        translated_runs = tagged_text_to_runs(initial_tagged_text)
        missing_ids = required_ids - set(translated_runs)
        if missing_ids:
            log(f"批量结果缺少 run 标签 {sorted(missing_ids)}，自动逐段重试。")
        else:
            apply_runs_translation(paragraph, translated_runs)
            if not has_cjk(paragraph_text(paragraph)):
                return True
            log("检测到译文仍有中文残留，自动逐段重试。")
        restore_run_texts(paragraph, original_run_texts)

    best_runs = {}
    for attempt in range(1, 3):
        translated_runs = translate_paragraph(client, provider, model, memory, previous_text, current, next_text, paragraph)
        best_runs = translated_runs
        missing_ids = required_ids - set(translated_runs)
        if missing_ids:
            log(f"逐段重试第 {attempt} 次仍缺少 run 标签 {sorted(missing_ids)}。")
        else:
            apply_runs_translation(paragraph, translated_runs)
            if not has_cjk(paragraph_text(paragraph)):
                return True
            log(f"逐段重试第 {attempt} 次后仍检测到中文残留。")
        restore_run_texts(paragraph, original_run_texts)

    fallback_ids = sorted((required_ids - set(best_runs)) | {idx for idx, text in best_runs.items() if has_cjk(text)})
    if fallback_ids:
        log(f"为保持格式不变，逐 run 补译：{fallback_ids}")
    for run_id in fallback_ids:
        original = original_run_texts[run_id]
        if not has_cjk(original):
            best_runs[run_id] = original
            continue
        best_runs[run_id] = translate_run_text(client, provider, model, memory, previous_text, current, next_text, original, run_id)
    for run_id in required_ids:
        best_runs.setdefault(run_id, original_run_texts[run_id])

    apply_runs_translation(paragraph, best_runs)
    if has_cjk(paragraph_text(paragraph)):
        log("逐 run 补译后仍有中文，请人工复核该段；格式已保持在原 run 上。")
        return False
    return True


def build_tagged_batches(paragraphs, targets: list[int], max_chars: int = 14000, max_segments: int = 24) -> list[list[dict]]:
    batches = []
    current_batch = []
    current_chars = 0
    for paragraph_index in targets:
        tagged_text = paragraph_to_tagged_text(paragraphs[paragraph_index])
        if not tagged_text:
            continue
        segment = {"id": str(paragraph_index), "tagged_text": tagged_text}
        segment_chars = len(tagged_text)
        if current_batch and (current_chars + segment_chars > max_chars or len(current_batch) >= max_segments):
            batches.append(current_batch)
            current_batch = []
            current_chars = 0
        current_batch.append(segment)
        current_chars += segment_chars
    if current_batch:
        batches.append(current_batch)
    return batches


def translate_tagged_batch(client: OpenAI, provider: str, model: str, memory: str, context_before: str, context_after: str, segments: list[dict]) -> dict[str, str]:
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate these Chinese legal contract segments into precise formal legal English.\n"
                "The text uses inline DOCX formatting tags like <r id=\"0\">text</r>. These tags represent Word formatting runs.\n"
                "Rules:\n"
                "1. Return JSON exactly like {\"segments\":[{\"id\":\"...\",\"tagged_text\":\"...\"}]}.\n"
                "2. Return every segment id exactly once.\n"
                "3. Preserve every <r id=\"...\"> and </r> tag exactly once per segment. Do not invent, delete, or renumber tags.\n"
                "4. Translate only the text between tags. Use surrounding segments and the document memory for legal context and terminology consistency.\n"
                "5. Keep numbers, dates, parties, article references, defined terms, and placeholders accurate.\n"
                "6. Translate all Chinese text, including bracketed placeholders and draft notes. Do not leave Chinese characters in the output.\n"
                "7. Do not output comments, Markdown, or explanations.\n\n"
                f"DOCUMENT MEMORY:\n{memory or '(none)'}\n\n"
                f"CONTEXT BEFORE THIS BATCH:\n{compact_text(context_before, 2000)}\n\n"
                f"CONTEXT AFTER THIS BATCH:\n{compact_text(context_after, 2000)}\n\n"
                f"SEGMENTS JSON:\n{json.dumps({'segments': segments}, ensure_ascii=False)}"
            ),
        },
    ]
    data = chat_json(client, provider, model, messages)
    translated = {}
    for item in data.get("segments", []):
        segment_id = str(item.get("id", "")).strip()
        tagged_text = str(item.get("tagged_text", ""))
        if segment_id:
            translated[segment_id] = tagged_text
    return translated


def available_fallback_path(output_path: Path) -> Path:
    stamp = time.strftime("%Y%m%d_%H%M%S")
    candidate = output_path.with_name(f"{output_path.stem}_自动另存_{stamp}{output_path.suffix}")
    counter = 1
    while candidate.exists():
        candidate = output_path.with_name(f"{output_path.stem}_自动另存_{stamp}_{counter}{output_path.suffix}")
        counter += 1
    return candidate


def safe_save_document(doc: Document, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f"{output_path.stem}.tmp.{int(time.time() * 1000)}{output_path.suffix}")
    doc.save(str(tmp_path))
    try:
        os.replace(str(tmp_path), str(output_path))
        return output_path
    except OSError as exc:
        if getattr(exc, "winerror", None) != 5 and getattr(exc, "errno", None) != 13:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass
            raise
        fallback_path = available_fallback_path(output_path)
        os.replace(str(tmp_path), str(fallback_path))
        return fallback_path


def translate_docx_run_concurrent(
    doc: Document,
    paragraphs: list,
    texts: list[str],
    targets: list[int],
    output_path: Path,
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    log,
    progress,
    save_event: threading.Event | None = None,
    autosave: bool = True,
) -> Path:
    run_texts_by_paragraph = [get_run_texts(paragraph) for paragraph in paragraphs]
    completed = 0
    total = len(targets)
    save_every = 10

    log(f"使用并发 run 保格式模式：{total} 个段落/单元，最多 {PARALLEL_WORKERS} 个 API 并发。")
    log("该模式不会合并或删除 Word run，只替换每个 run 的文字，因此粗体、下划线、高亮会留在原 run 上。")

    with ThreadPoolExecutor(max_workers=PARALLEL_WORKERS) as executor:
        future_map = {
            executor.submit(
                translate_paragraph_task,
                client,
                provider,
                model,
                memory,
                texts,
                run_texts_by_paragraph,
                paragraph_index,
            ): paragraph_index
            for paragraph_index in targets
        }

        for future in as_completed(future_map):
            paragraph_index = future_map[future]
            try:
                paragraph_index, translated_runs, warnings = future.result()
            except Exception as exc:
                log(f"段落 {paragraph_index + 1} 翻译失败，将保留原文并继续。原因：{exc}")
                translated_runs = {}
                warnings = []

            if translated_runs:
                apply_runs_translation(paragraphs[paragraph_index], translated_runs)
                if has_cjk(paragraph_text(paragraphs[paragraph_index])):
                    log(f"段落 {paragraph_index + 1} 翻译后仍检测到中文，请人工复核。")
            for warning in warnings:
                log(f"段落 {paragraph_index + 1}：{warning}")

            completed += 1
            progress(completed, total, compact_text(texts[paragraph_index], 120))

            requested_save = save_event.is_set() if save_event else False
            if (autosave and (completed % save_every == 0 or completed == total)) or requested_save:
                try:
                    saved_path = safe_save_document(doc, output_path)
                    if saved_path != output_path:
                        log(f"输出文件被占用，已自动另存为：{saved_path}")
                        output_path = saved_path
                    log(f"已保存当前进度：{output_path}（{completed}/{total}，{completed / total:.0%}）")
                except Exception as exc:
                    log(f"保存当前进度失败：{exc}")
                finally:
                    if save_event:
                        save_event.clear()

    saved_path = safe_save_document(doc, output_path)
    if saved_path != output_path:
        log(f"输出文件被占用，已自动另存为：{saved_path}")
        output_path = saved_path
    progress(total, total, "全部翻译完成")
    log(f"完成：{output_path}")
    return output_path


def translate_docx(
    input_path: Path,
    output_path: Path,
    provider: str,
    api_key: str,
    base_url: str,
    model: str,
    log,
    progress,
    save_event: threading.Event | None = None,
    autosave: bool = True,
    translation_mode: str = MODE_RUN_CONCURRENT,
) -> None:
    doc = Document(str(input_path))
    paragraphs = list(iter_document_paragraphs(doc))
    texts = [paragraph_text(p) for p in paragraphs]
    targets = [i for i, text in enumerate(texts) if has_cjk(text)]

    if not targets:
        raise ValueError("没有在文档中发现中文内容。")

    client = build_client(api_key, base_url)
    log(f"使用 API：{provider} / {model}")
    if base_url.strip():
        log(f"Base URL：{base_url.strip()}")
    progress(0, len(targets), "正在读取整份合同并提取术语表...")
    memory = make_document_memory(client, provider, model, texts, log)
    log(f"发现 {len(targets)} 个含中文的段落/表格单元，开始翻译...")

    if translation_mode == MODE_RUN_CONCURRENT:
        return translate_docx_run_concurrent(
            doc,
            paragraphs,
            texts,
            targets,
            output_path,
            client,
            provider,
            model,
            memory,
            log,
            progress,
            save_event,
            autosave,
        )

    if translation_mode == MODE_FAST:
        batches = build_tagged_batches(paragraphs, targets)
        log(f"使用方案二：批量标签模式。将 {len(targets)} 个段落/单元合并为 {len(batches)} 次 API 调用。")
        completed = 0
        for batch_number, batch in enumerate(batches, start=1):
            first_index = int(batch[0]["id"])
            last_index = int(batch[-1]["id"])
            context_before = "\n".join(texts[max(0, first_index - 3) : first_index])
            context_after = "\n".join(texts[last_index + 1 : min(len(texts), last_index + 4)])
            batch_preview = compact_text(paragraph_text(paragraphs[first_index]), 120)

            progress(completed, len(targets), f"批量翻译第 {batch_number}/{len(batches)} 批：{batch_preview}")
            log(f"[批次 {batch_number}/{len(batches)}] 正在翻译 {len(batch)} 个段落/单元...")

            try:
                translated_batch = translate_tagged_batch(client, provider, model, memory, context_before, context_after, batch)
            except Exception as exc:
                log(f"批量翻译失败，改用逐段模式处理本批次。原因：{exc}")
                translated_batch = {}

            for segment in batch:
                paragraph_index = int(segment["id"])
                translated_text = translated_batch.get(segment["id"], "")
                translate_paragraph_with_fallback(
                    client,
                    provider,
                    model,
                    memory,
                    texts,
                    paragraphs,
                    paragraph_index,
                    log,
                    translated_text or None,
                )
                completed += 1
                current_preview = compact_text(texts[paragraph_index], 120)
                progress(completed, len(targets), current_preview)

            requested_save = save_event.is_set() if save_event else False
            if autosave or requested_save:
                try:
                    saved_path = safe_save_document(doc, output_path)
                    if saved_path != output_path:
                        log(f"输出文件被占用，已自动另存为：{saved_path}")
                        output_path = saved_path
                    log(f"已保存当前进度：{output_path}（{completed}/{len(targets)}，{completed / len(targets):.0%}）")
                except Exception as exc:
                    log(f"保存当前进度失败：{exc}")
                finally:
                    if save_event:
                        save_event.clear()

        saved_path = safe_save_document(doc, output_path)
        if saved_path != output_path:
            log(f"输出文件被占用，已自动另存为：{saved_path}")
            output_path = saved_path
        progress(len(targets), len(targets), "全部翻译完成")
        log(f"完成：{output_path}")
        return output_path

    for count, paragraph_index in enumerate(targets, start=1):
        paragraph = paragraphs[paragraph_index]
        current = texts[paragraph_index]
        previous_text = texts[paragraph_index - 1] if paragraph_index > 0 else ""
        next_text = texts[paragraph_index + 1] if paragraph_index + 1 < len(texts) else ""
        current_preview = compact_text(current, 120)

        progress(count - 1, len(targets), current_preview)
        log(f"[{count}/{len(targets)}] 正在翻译：{compact_text(current, 70)}")
        translate_paragraph_with_fallback(
            client,
            provider,
            model,
            memory,
            texts,
            paragraphs,
            paragraph_index,
            log,
        )

        progress(count, len(targets), current_preview)
        requested_save = save_event.is_set() if save_event else False
        if autosave or requested_save:
            try:
                saved_path = safe_save_document(doc, output_path)
                if saved_path != output_path:
                    log(f"输出文件被占用，已自动另存为：{saved_path}")
                    output_path = saved_path
                log(f"已保存当前进度：{output_path}（{count}/{len(targets)}，{count / len(targets):.0%}）")
            except Exception as exc:
                log(f"保存当前进度失败：{exc}")
            finally:
                if save_event:
                    save_event.clear()

    saved_path = safe_save_document(doc, output_path)
    if saved_path != output_path:
        log(f"输出文件被占用，已自动另存为：{saved_path}")
        output_path = saved_path
    progress(len(targets), len(targets), "全部翻译完成")
    log(f"完成：{output_path}")
    return output_path


class TranslatorApp:
    def __init__(self, root):
        self.root = root
        self.main_thread_id = threading.get_ident()
        self.root.title("法律合同中译英工具 - OpenAI / DeepSeek")
        self.root.geometry("840x740")
        self.root.minsize(780, 680)

        config = load_config()
        saved_provider = config.get("provider", "OpenAI")
        if saved_provider not in PROVIDER_DEFAULTS:
            saved_provider = "OpenAI"
        provider_defaults = PROVIDER_DEFAULTS[saved_provider]
        self.provider = StringVar(value=saved_provider)
        self.api_key = StringVar(value=config.get("api_key", ""))
        self.base_url = StringVar(value=config.get("base_url", provider_defaults["base_url"]))
        self.model = StringVar(value=config.get("model", provider_defaults["model"]))
        saved_mode = config.get("translation_mode", MODE_RUN_CONCURRENT)
        if saved_mode not in (MODE_RUN_CONCURRENT, MODE_FAST, MODE_CAREFUL):
            saved_mode = MODE_RUN_CONCURRENT
        self.translation_mode = StringVar(value=saved_mode)
        self.file_path = StringVar(value="")
        self.output_path = StringVar(value="")
        self.status = StringVar(value="请选择或拖入 .docx 文件")
        self.progress_value = DoubleVar(value=0)
        self.progress_text = StringVar(value="进度：0%")
        self.current_text = StringVar(value="当前段落：尚未开始")
        self.remember_key = BooleanVar(value=bool(config.get("api_key")))
        self.autosave_progress = BooleanVar(value=True)
        self.save_event = threading.Event()

        top = Frame(root, padx=16, pady=12)
        top.pack(fill=X)

        Label(top, text="API 类型").pack(anchor="w")
        OptionMenu(top, self.provider, *PROVIDER_DEFAULTS.keys(), command=self.on_provider_change).pack(fill=X, pady=(4, 8))

        self.api_key_label = Label(top, text=provider_defaults["key_label"])
        self.api_key_label.pack(anchor="w")
        Entry(top, textvariable=self.api_key, show="*", width=80).pack(fill=X, pady=(4, 8))

        Label(top, text="Base URL").pack(anchor="w")
        Entry(top, textvariable=self.base_url, width=80).pack(fill=X, pady=(4, 8))

        Label(top, text="模型名称").pack(anchor="w")
        Entry(top, textvariable=self.model, width=80).pack(fill=X, pady=(4, 8))

        Label(top, text="翻译模式").pack(anchor="w")
        OptionMenu(top, self.translation_mode, MODE_RUN_CONCURRENT, MODE_FAST, MODE_CAREFUL).pack(fill=X, pady=(4, 8))

        Checkbutton(top, text="记住 API Key（明文保存在本机此工具目录）", variable=self.remember_key).pack(anchor="w")

        file_frame = Frame(root, padx=16, pady=8)
        file_frame.pack(fill=X)
        Label(file_frame, text="合同文件（仅支持 .docx）").pack(anchor="w")
        row = Frame(file_frame)
        row.pack(fill=X, pady=(4, 0))
        Entry(row, textvariable=self.file_path).pack(side=LEFT, fill=X, expand=True)
        Button(row, text="选择文件", command=self.choose_file, width=12).pack(side=RIGHT, padx=(8, 0))

        Label(file_frame, text="输出文件").pack(anchor="w", pady=(10, 0))
        output_row = Frame(file_frame)
        output_row.pack(fill=X, pady=(4, 0))
        Entry(output_row, textvariable=self.output_path).pack(side=LEFT, fill=X, expand=True)
        Button(output_row, text="另存为", command=self.choose_output_file, width=12).pack(side=RIGHT, padx=(8, 0))

        self.drop_label = Label(
            root,
            text="把 Word 合同拖到这里\n支持保留正文、表格、页眉页脚里的粗体 / 下划线 / 高亮等 run 格式",
            relief="groove",
            height=5,
            padx=12,
            pady=12,
        )
        self.drop_label.pack(fill=X, padx=16, pady=8)
        if DND_AVAILABLE:
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind("<<Drop>>", self.on_drop)
        else:
            self.drop_label.config(text=self.drop_label.cget("text") + "\n\n提示：拖拽组件未加载，可用“选择文件”。")

        progress_frame = Frame(root, padx=16, pady=8)
        progress_frame.pack(fill=X)
        Label(progress_frame, textvariable=self.progress_text).pack(anchor="w")
        Progressbar(progress_frame, variable=self.progress_value, maximum=100).pack(fill=X, pady=(4, 6))
        Label(progress_frame, textvariable=self.current_text, wraplength=760, justify=LEFT).pack(anchor="w")

        action_frame = Frame(root, padx=16, pady=8)
        action_frame.pack(fill=X)
        Checkbutton(action_frame, text="自动把已翻译进度保存到输出文件", variable=self.autosave_progress).pack(anchor="w", pady=(0, 8))
        buttons = Frame(action_frame)
        buttons.pack(fill=X)
        self.start_button = Button(buttons, text="开始翻译并导出", command=self.start_translation, height=2)
        self.start_button.pack(side=LEFT, fill=X, expand=True)
        self.save_button = Button(buttons, text="保存当前进度", command=self.save_current_progress, height=2, state="disabled")
        self.save_button.pack(side=RIGHT, padx=(8, 0))


        Label(root, textvariable=self.status, padx=16).pack(fill=X, anchor="w")
        self.log_box = Text(root, height=14, wrap="word")
        self.log_box.pack(fill=BOTH, expand=True, padx=16, pady=(4, 16))

        self.log("使用方式：输入 API Key，拖入或选择 .docx 文件，然后点击“开始翻译并导出”。")
        self.log("API 类型可选 OpenAI 或 DeepSeek；DeepSeek 默认使用 https://api.deepseek.com 和 deepseek-v4-flash。")
        self.log("默认使用并发 run 保格式模式：保留粗体、下划线、高亮等原 Word run 格式，并用并发提高速度。")
        self.log("重要提示：法律文件请务必由专业律师或法律译审复核；本工具不是法律意见。")

    def on_provider_change(self, selected_provider):
        defaults = PROVIDER_DEFAULTS.get(selected_provider, PROVIDER_DEFAULTS["OpenAI"])
        self.api_key_label.config(text=defaults["key_label"])
        self.base_url.set(defaults["base_url"])
        self.model.set(defaults["model"])

    def default_output_for(self, input_path: Path) -> Path:
        return input_path.with_name(input_path.stem + "_英文翻译.docx")

    def set_input_file(self, path: str, status: str):
        input_path = Path(path)
        self.file_path.set(str(input_path))
        self.output_path.set(str(self.default_output_for(input_path)))
        self.status.set(status)

    def choose_file(self):
        path = filedialog.askopenfilename(title="选择 Word 合同", filetypes=[("Word 文档", "*.docx")])
        if path:
            self.set_input_file(path, "已选择文件")

    def choose_output_file(self):
        input_text = self.file_path.get().strip().strip('"')
        initialdir = APP_DIR
        initialfile = "英文翻译.docx"
        if input_text:
            input_path = Path(input_text)
            initialdir = input_path.parent if input_path.parent.exists() else APP_DIR
            initialfile = self.default_output_for(input_path).name
        path = filedialog.asksaveasfilename(
            title="选择输出 Word 文件",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialdir=str(initialdir),
            initialfile=initialfile,
        )
        if path:
            output_path = Path(path)
            if output_path.suffix.lower() != ".docx":
                output_path = output_path.with_suffix(".docx")
            self.output_path.set(str(output_path))
            self.status.set("已选择输出文件")

    def on_drop(self, event):
        paths = self.root.tk.splitlist(event.data)
        if paths:
            self.set_input_file(paths[0], "已拖入文件")

    def log(self, text: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.log, text)
            return
        self.log_box.insert(END, text + "\n")
        self.log_box.see(END)
        self.root.update_idletasks()

    def update_progress(self, done: int, total: int, current: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.update_progress, done, total, current)
            return
        if total <= 0:
            percent = 0
        else:
            percent = max(0, min(100, done * 100 / total))
        self.progress_value.set(percent)
        self.progress_text.set(f"进度：{done}/{total}（{percent:.1f}%）")
        self.status.set(f"翻译中：{done}/{total}（{percent:.1f}%）")
        self.current_text.set(f"当前段落：{current}")
        self.root.update_idletasks()

    def save_current_progress(self):
        self.save_event.set()
        self.log("已请求保存当前进度；会在当前段落翻译完成后写入输出文件。")
        self.status.set("正在等待当前段落完成后保存进度...")

    def start_translation(self):
        provider = self.provider.get().strip()
        translation_mode = self.translation_mode.get().strip() or MODE_RUN_CONCURRENT
        api_key = self.api_key.get().strip()
        base_url = self.base_url.get().strip()
        model = self.model.get().strip()
        input_text = self.file_path.get().strip().strip('"')
        output_text = self.output_path.get().strip().strip('"')

        if not api_key:
            messagebox.showerror("缺少 API Key", f"请先输入 {provider} API Key。")
            return
        if not model:
            messagebox.showerror("缺少模型", "请填写模型名称。")
            return
        if not input_text:
            messagebox.showerror("缺少文件", "请拖入或选择一个 .docx 文件。")
            return

        input_path = Path(input_text)
        if input_path.suffix.lower() != ".docx" or not input_path.exists():
            messagebox.showerror("文件不支持", "目前只支持存在的 .docx 文件。请先把 PDF 或 WPS 文件另存为 Word .docx。")
            return

        if not output_text:
            output_path = self.default_output_for(input_path)
            self.output_path.set(str(output_path))
        else:
            output_path = Path(output_text)
            if output_path.suffix.lower() != ".docx":
                output_path = output_path.with_suffix(".docx")
                self.output_path.set(str(output_path))

        if not output_path.parent.exists():
            messagebox.showerror("输出路径不存在", "请选择一个已经存在的文件夹作为输出位置。")
            return
        if output_path.resolve() == input_path.resolve():
            messagebox.showerror("输出文件不能覆盖原文", "请把输出文件保存为另一个 .docx 文件。")
            return

        if provider == "DeepSeek" and not base_url:
            base_url = PROVIDER_DEFAULTS["DeepSeek"]["base_url"]
            self.base_url.set(base_url)

        if self.remember_key.get():
            save_config(provider, api_key, base_url, model, translation_mode)

        self.save_event.clear()
        self.progress_value.set(0)
        self.progress_text.set("进度：0%")
        self.current_text.set("当前段落：准备开始")
        self.start_button.config(state="disabled")
        self.save_button.config(state="normal")
        self.status.set("翻译中，请不要关闭窗口")
        thread = threading.Thread(
            target=self.worker,
            args=(input_path, output_path, provider, api_key, base_url, model, self.autosave_progress.get(), translation_mode),
            daemon=True,
        )
        thread.start()

    def worker(self, input_path: Path, output_path: Path, provider: str, api_key: str, base_url: str, model: str, autosave: bool, translation_mode: str):
        try:
            actual_output_path = translate_docx(
                input_path,
                output_path,
                provider,
                api_key,
                base_url,
                model,
                self.log,
                self.update_progress,
                self.save_event,
                autosave,
                translation_mode,
            )
            self.root.after(0, self.finish_success, actual_output_path or output_path)
        except Exception as exc:
            self.log("发生错误：")
            self.log(str(exc))
            self.log(traceback.format_exc())
            self.root.after(0, self.finish_error, str(exc))

    def finish_success(self, output_path: Path):
        self.status.set("翻译完成")
        self.start_button.config(state="normal")
        self.save_button.config(state="disabled")
        self.save_event.clear()
        messagebox.showinfo("完成", f"已导出：\n{output_path}")

    def finish_error(self, error_message: str):
        self.status.set("翻译失败")
        self.start_button.config(state="normal")
        self.save_button.config(state="disabled")
        self.save_event.clear()
        messagebox.showerror("翻译失败", error_message)


def main():
    root_class = TkinterDnD.Tk if DND_AVAILABLE else Tk
    root = root_class()
    TranslatorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
