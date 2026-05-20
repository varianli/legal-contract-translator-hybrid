import json
import os
import re
import threading
import time
import traceback
from copy import copy
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from tkinter import BOTH, END, LEFT, RIGHT, X, BooleanVar, Button, Checkbutton, Entry, Frame, Label, OptionMenu, StringVar, Text, Tk, filedialog, messagebox
from tkinter.ttk import Progressbar

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD

    DND_AVAILABLE = True
except Exception:
    TkinterDnD = None
    DND_FILES = None
    DND_AVAILABLE = False


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "输出"
CONFIG_PATH = APP_DIR / ".new_scheme_config.json"

PROVIDER_DEFAULTS = {
    "OpenAI": {"base_url": "", "model": "gpt-4.1", "key_label": "OpenAI API Key"},
    "DeepSeek": {"base_url": "https://api.deepseek.com", "model": "deepseek-v4-flash", "key_label": "DeepSeek API Key"},
}
APP_VERSION = "v1.4"

SYSTEM_PROMPT = """You are a senior bilingual legal translator and legal formatting alignment specialist.
Translate Chinese legal contracts into precise, formal legal English.
Preserve legal meaning, defined terms, numbering, dates, parties, amounts, placeholders, and clause references.
Return only valid JSON when JSON is requested."""

CAPITAL_MARKETS_LEGAL_RAG = """Capital markets legal English retrieval notes:
- This is a capital markets / private equity style legal contract. Prefer formal transactional drafting, not conversational English.
- Translate \u9644\u5f55 as Appendix and keep it consistent. Do not translate \u9644\u5f55 as Schedule.
- Translate \u9644\u4ef6 as Annex. Translate \u9644\u8868 as Schedule.
- Use Roman numerals for Chinese appendix numerals: \u9644\u5f55\u4e00 -> Appendix I; \u9644\u5f55\u4e8c -> Appendix II; \u9644\u5f55\u4e09\uff08A\uff09 -> Appendix III(A); \u9644\u5f55\u4e03 -> Appendix VII.
- Translate \u62ab\u9732\u51fd / \u62ab\u9732\u6e05\u5355 as Disclosure Schedule. If it is titled \u9644\u5f55\u4e94\uff08\u62ab\u9732\u51fd\uff09, use Appendix V (Disclosure Schedule).
- Translate \u4ea4\u5272 as Closing, \u4ea4\u5272\u65e5 as Closing Date, \u6700\u8fdf\u5b8c\u6210\u65e5 as Longstop Date, \u7b7e\u7f72\u65e5 as Signing Date.
- Translate \u672c\u6b21\u4ea4\u6613 as this Transaction, \u6295\u8d44\u6b3e as Investment Amount, \u76ee\u6807\u516c\u53f8 as Target Company, \u96c6\u56e2\u516c\u53f8 as Group Company / Group Companies.
- For RMB amounts expressed as \u4eba\u6c11\u5e01X\u4e07\u5143, convert to yuan in English: \u4eba\u6c11\u5e0112.7764\u4e07\u5143 -> RMB 127,764. Do not output Ten Thousand Yuan or million for \u4e07\u5143.
- Preserve source \u3010...\u3011 bracket placeholders exactly. \u3010\u3011 stays \u3010\u3011, not [] or underscores; \u3010number\u3011 stays in \u3010number\u3011 form after translation/conversion.
- Preserve defined-term capitalization once established. If the same Chinese term recurs, reuse the same English term.
- Use English punctuation in English output: straight quotes, half-width commas, periods, semicolons, colons, and parentheses, while preserving source \u3010...\u3011 placeholders."""

LEGAL_RAG_TERMS = [
    ("\u9644\u5f55", "\u9644\u5f55 = Appendix; use Roman numerals, e.g. Appendix I, Appendix II, Appendix III(A)."),
    ("\u9644\u4ef6", "\u9644\u4ef6 = Annex."),
    ("\u9644\u8868", "\u9644\u8868 = Schedule."),
    ("\u62ab\u9732\u51fd", "\u62ab\u9732\u51fd = Disclosure Schedule."),
    ("\u4ea4\u5272\u65e5", "\u4ea4\u5272\u65e5 = Closing Date."),
    ("\u4ea4\u5272", "\u4ea4\u5272 = Closing."),
    ("\u6700\u8fdf\u5b8c\u6210\u65e5", "\u6700\u8fdf\u5b8c\u6210\u65e5 = Longstop Date."),
    ("\u7b7e\u7f72\u65e5", "\u7b7e\u7f72\u65e5 = Signing Date."),
    ("\u672c\u6b21\u4ea4\u6613", "\u672c\u6b21\u4ea4\u6613 = this Transaction."),
    ("\u6295\u8d44\u6b3e", "\u6295\u8d44\u6b3e = Investment Amount."),
    ("\u76ee\u6807\u516c\u53f8", "\u76ee\u6807\u516c\u53f8 = Target Company."),
    ("\u96c6\u56e2\u516c\u53f8", "\u96c6\u56e2\u516c\u53f8 = Group Company / Group Companies."),
]

CHINESE_NUMERAL_ROMAN = {
    "\u4e00": "I",
    "\u4e8c": "II",
    "\u4e09": "III",
    "\u56db": "IV",
    "\u4e94": "V",
    "\u516d": "VI",
    "\u4e03": "VII",
    "\u516b": "VIII",
    "\u4e5d": "IX",
    "\u5341": "X",
}
ARABIC_ROMAN = {str(index): roman for index, roman in enumerate(["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]) if index}
ROMAN_ARABIC = {roman: arabic for arabic, roman in ARABIC_ROMAN.items()}
APPENDIX_SOURCE_RE = re.compile(r"(\u9644\u5f55|\u9644\u4ef6|\u9644\u8868)([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+)(?:[\uff08(]\s*([A-Za-z\uff21-\uff3a])\s*[\uff09)])?")
SOURCE_FALLBACK_TARGETS = {
    "\u5b9a\u4e49": ["Definitions"],
    "\u672c\u6b21\u4ea4\u6613\u5b89\u6392": ["Transaction Arrangement", "Transaction Arrangements"],
    "\u4ea4\u5272": ["Closing"],
    "\u4ea4\u5272\u540e\u4e49\u52a1": ["Post-Closing Obligations", "Post-closing Obligations"],
    "\u58f0\u660e\u548c\u4fdd\u8bc1": ["Representations and Warranties"],
}

CJK_PUNCT_TRANSLATION = str.maketrans(
    {
        "\uff0c": ",",
        "\u3002": ".",
        "\uff1b": ";",
        "\uff1a": ":",
        "\uff01": "!",
        "\uff1f": "?",
        "\uff08": "(",
        "\uff09": ")",
        "\u201c": '"',
        "\u201d": '"',
        "\u2018": "'",
        "\u2019": "'",
        "\u3001": ",",
        "\u300a": '"',
        "\u300b": '"',
    }
)


@dataclass
class FormatSpan:
    span_id: str
    source_text: str
    style_label: str
    bold: bool
    italic: bool
    underline: object
    highlight_color: object


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
W_VAL = f"{W_NS}val"
OFF_VALUES = {"0", "false", "off", "none"}


def has_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text or ""))


def compact_text(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if len(text) <= limit:
        return text
    half = max(1, limit // 2)
    return text[:half] + "\n...[middle omitted]...\n" + text[-half:]


def legal_rag_for_text(text: str) -> str:
    retrieved = []
    text = text or ""
    for trigger, note in LEGAL_RAG_TERMS:
        if trigger in text and note not in retrieved:
            retrieved.append(note)
    if not retrieved:
        return CAPITAL_MARKETS_LEGAL_RAG
    return CAPITAL_MARKETS_LEGAL_RAG + "\n\nRetrieved terms for this chunk:\n- " + "\n- ".join(retrieved)


def normalize_english_punctuation(text: str) -> str:
    text = (text or "").translate(CJK_PUNCT_TRANSLATION)
    text = re.sub(r"\s+([,.;:!?\]\)])", r"\1", text)
    text = re.sub(r"([\[\(])\s+", r"\1", text)
    text = re.sub(r"([,;:!?])(?=[A-Za-z\"'])", r"\1 ", text)
    text = re.sub(r"(?<!\d)(\.)(?=[A-Za-z\"'])", r"\1 ", text)
    text = re.sub(r"(?<=\d),\s+(?=\d{3}\b)", ",", text)
    text = re.sub(r"(?<=\d)\.\s+(?=\d)", ".", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


FULL_RMB_WANYUAN_RE = re.compile(r"^\s*\u4eba\u6c11\u5e01\s*([\u3010\[])?\s*([0-9][0-9,，]*(?:\.\d+)?)\s*([\u3011\]])?\s*\u4e07\u5143\s*$")
FULL_CHINESE_DATE_RE = re.compile(r"^\s*(\d{4})\s*\u5e74\s*(\d{1,2}|\u3010\u3011|_+)\s*\u6708\s*(\d{1,2}|\u3010\u3011|_+)\s*\u65e5\s*$")
SOURCE_BRACKET_RE = re.compile(r"\u3010([^\u3011]*)\u3011")
MONTH_NAMES = {
    1: "January",
    2: "February",
    3: "March",
    4: "April",
    5: "May",
    6: "June",
    7: "July",
    8: "August",
    9: "September",
    10: "October",
    11: "November",
    12: "December",
}


def format_rmb_yuan_from_wanyuan(number_text: str) -> str | None:
    cleaned = (number_text or "").replace(",", "").replace("，", "").strip()
    try:
        amount = (Decimal(cleaned) * Decimal("10000")).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError):
        return None
    return f"{int(amount):,}"


def deterministic_amount_translation(source_text: str) -> str | None:
    match = FULL_RMB_WANYUAN_RE.fullmatch(source_text or "")
    if not match:
        return None
    amount = format_rmb_yuan_from_wanyuan(match.group(2))
    if not amount:
        return None
    if match.group(1) or match.group(3):
        return f"RMB \u3010{amount}\u3011"
    return f"RMB {amount}"


def date_token_to_english(token: str, is_month: bool) -> str:
    token = (token or "").strip()
    if token == "\u3010\u3011" or re.fullmatch(r"_+", token):
        return token
    if not token.isdigit():
        return token
    value = int(token)
    if is_month and value in MONTH_NAMES:
        return MONTH_NAMES[value]
    return str(value)


def deterministic_date_translation(source_text: str) -> str | None:
    match = FULL_CHINESE_DATE_RE.fullmatch(source_text or "")
    if not match:
        return None
    year, month_token, day_token = match.groups()
    month = date_token_to_english(month_token, True)
    day = date_token_to_english(day_token, False)
    return f"{month} {day}, {year}"


def preserve_source_bracket_tokens(source_text: str, translation: str) -> str:
    result = translation or ""
    for content in SOURCE_BRACKET_RE.findall(source_text or ""):
        target = f"\u3010{content}\u3011"
        if target in result:
            continue
        if content:
            result = re.sub(r"\[\s*" + re.escape(content) + r"\s*\]", target, result, count=1)
        else:
            result = re.sub(r"\[\s*\]", target, result, count=1)
            result = re.sub(r"_+", target, result, count=1)
    return result


def normalize_translation_against_source(source_text: str, translation: str) -> str:
    deterministic = deterministic_amount_translation(source_text)
    if deterministic:
        return deterministic
    deterministic = deterministic_date_translation(source_text)
    if deterministic:
        return deterministic
    return preserve_source_bracket_tokens(source_text, normalize_legal_english(translation))


def normalize_roman_token(token: str) -> str:
    token = (token or "").strip().upper()
    if token in CHINESE_NUMERAL_ROMAN:
        return CHINESE_NUMERAL_ROMAN[token]
    if token in ARABIC_ROMAN:
        return ARABIC_ROMAN[token]
    if token in ROMAN_ARABIC:
        return token
    return token


def normalize_legal_english(text: str) -> str:
    text = normalize_english_punctuation(text)

    def appendix_repl(match):
        roman = normalize_roman_token(match.group(1))
        suffix = f"({match.group(2).upper()})" if match.group(2) else ""
        return f"Appendix {roman}{suffix}"

    text = re.sub(r"\bAppendix\s+([1-9]|10|I|II|III|IV|V|VI|VII|VIII|IX|X)(?!\.\d)(?:\s*\(\s*([A-Za-z])\s*\))?", appendix_repl, text, flags=re.IGNORECASE)
    text = re.sub(r"\bSchedule\s+([1-9]|10|I|II|III|IV|V|VI|VII|VIII|IX|X)\s*\(?\s*Disclosure Schedule\s*\)?", lambda m: f"Appendix {normalize_roman_token(m.group(1))} (Disclosure Schedule)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bSchedule\s+([1-9]|10|I|II|III|IV|V|VI|VII|VIII|IX|X)\s+of\s+the\s+Disclosure\s+Schedule\b", lambda m: f"Appendix {normalize_roman_token(m.group(1))} (Disclosure Schedule)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bLong\s+Stop\s+Date\b", "Longstop Date", text, flags=re.IGNORECASE)
    return text


def source_span_fallback_targets(source_text: str) -> list[str]:
    source_text = source_text or ""
    direct_targets = SOURCE_FALLBACK_TARGETS.get(source_text.strip(), [])
    if direct_targets:
        return direct_targets
    match = APPENDIX_SOURCE_RE.search(source_text or "")
    if not match:
        return []
    label_map = {"\u9644\u5f55": "Appendix", "\u9644\u4ef6": "Annex", "\u9644\u8868": "Schedule"}
    label = label_map.get(match.group(1))
    roman = normalize_roman_token(match.group(2))
    if not label or not roman:
        return []
    suffix = ""
    if match.group(3):
        suffix = f"({match.group(3).upper()})"
    targets = [f"{label} {roman}{suffix}"]
    arabic = ROMAN_ARABIC.get(roman)
    if arabic:
        targets.append(f"{label} {arabic}{suffix}")
    return targets


def load_config() -> dict:
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def save_config(provider: str, api_key: str, base_url: str, model: str) -> None:
    CONFIG_PATH.write_text(
        json.dumps({"provider": provider, "api_key": api_key, "base_url": base_url, "model": model}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def parse_json_object(raw: str) -> dict:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r"```$", "", raw).strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("模型没有返回 JSON 对象。")
    return json.loads(raw[start : end + 1])


def build_client(api_key: str, base_url: str) -> OpenAI:
    kwargs = {"api_key": api_key, "timeout": 180}
    if base_url.strip():
        kwargs["base_url"] = base_url.strip()
    return OpenAI(**kwargs)


def chat_json(client: OpenAI, provider: str, model: str, messages: list[dict], retries: int = 3) -> dict:
    last_error = None
    for attempt in range(retries):
        response_format_options = [True, False] if provider == "DeepSeek" else [True]
        for use_response_format in response_format_options:
            try:
                kwargs = {"model": model, "messages": messages, "temperature": 0.05}
                if use_response_format:
                    kwargs["response_format"] = {"type": "json_object"}
                response = client.chat.completions.create(**kwargs)
                return parse_json_object(response.choices[0].message.content)
            except Exception as exc:
                last_error = exc
        if attempt + 1 < retries:
            time.sleep(2 + attempt * 3)
    raise last_error


def iter_table_paragraphs(table, seen):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_cell_paragraphs(cell, seen)


def iter_cell_paragraphs(cell, seen):
    for paragraph in cell.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in cell.tables:
        yield from iter_table_paragraphs(table, seen)


def iter_part_paragraphs(part, seen):
    for paragraph in part.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in part.tables:
        yield from iter_table_paragraphs(table, seen)


def iter_document_paragraphs(doc: Document):
    seen = set()
    yield from iter_part_paragraphs(doc, seen)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            yield from iter_part_paragraphs(part, seen)


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def xml_has_off_property(element, tag_name: str) -> bool:
    if element is None:
        return False
    for child in element.iter(f"{W_NS}{tag_name}"):
        value = child.get(W_VAL)
        if value is not None and value.lower() in OFF_VALUES:
            return True
    return False


def xml_has_on_property(element, tag_name: str) -> bool:
    if element is None:
        return False
    for child in element.iter(f"{W_NS}{tag_name}"):
        value = child.get(W_VAL)
        if value is None or value.lower() not in OFF_VALUES:
            return True
    return False


def style_format_defaults(paragraph) -> tuple:
    style = getattr(paragraph, "style", None)
    font = getattr(style, "font", None)
    if font is None:
        return (False, False, None, None)
    return (bool(font.bold), bool(font.italic), font.underline, font.highlight_color)


def run_style_tuple(run, paragraph=None) -> tuple:
    direct = (bool(run.bold), bool(run.italic), run.underline, run.font.highlight_color)
    if paragraph is None:
        return direct
    style_bold, style_italic, style_underline, style_highlight = style_format_defaults(paragraph)
    rpr = run._r.rPr

    bold = True if run.bold is True or xml_has_on_property(rpr, "b") else False
    if not bold and run.bold is not False and not xml_has_off_property(rpr, "b") and style_bold:
        bold = True

    italic = True if run.italic is True or xml_has_on_property(rpr, "i") else False
    if not italic and run.italic is not False and not xml_has_off_property(rpr, "i") and style_italic:
        italic = True

    underline = run.underline
    if not underline and run.underline is not False and not xml_has_off_property(rpr, "u"):
        underline = style_underline if style_underline else (True if xml_has_on_property(rpr, "u") else None)

    highlight = run.font.highlight_color
    if not highlight and not xml_has_off_property(rpr, "highlight"):
        highlight = style_highlight
    return (bold, italic, underline, highlight)


def style_label(style: tuple) -> str:
    bold, italic, underline, highlight = style
    labels = []
    if bold:
        labels.append("bold")
    if italic:
        labels.append("italic")
    if underline:
        labels.append("underline")
    if highlight:
        labels.append(f"highlight:{highlight}")
    return "+".join(labels)


def is_formatted_run(run, paragraph=None) -> bool:
    bold, italic, underline, highlight = run_style_tuple(run, paragraph)
    return bool((bold or italic or underline or highlight) and run.text and run.text.strip())


def extract_format_spans(paragraph, paragraph_id: int) -> list[FormatSpan]:
    spans: list[FormatSpan] = []
    current_text = ""
    current_style = None
    span_number = 0

    def flush():
        nonlocal current_text, current_style, span_number
        if current_text.strip() and current_style:
            bold, italic, underline, highlight = current_style
            spans.append(
                FormatSpan(
                    span_id=f"p{paragraph_id}_s{span_number}",
                    source_text=current_text,
                    style_label=style_label(current_style),
                    bold=bold,
                    italic=italic,
                    underline=copy(underline),
                    highlight_color=copy(highlight),
                )
            )
            span_number += 1
        current_text = ""
        current_style = None

    for run in paragraph.runs:
        if is_formatted_run(run, paragraph):
            style = run_style_tuple(run, paragraph)
            if current_style == style:
                current_text += run.text
            else:
                flush()
                current_style = style
                current_text = run.text
        else:
            flush()
    flush()
    return spans


def make_translation_memory(client: OpenAI, provider: str, model: str, paragraphs: list[str], log) -> str:
    text = "\n".join(p for p in paragraphs if has_cjk(p))
    if not text.strip():
        return ""
    log("正在读取整份合同，提取术语和翻译规则...")
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Create a concise bilingual translation memory for this legal contract.\n"
                "Return JSON: {\"memory\":\"...\"}.\n"
                "Include parties, defined terms, recurring legal phrases, placeholders, draft-note style, and translation preferences.\n"
                "Keep it concise but useful for consistent contract translation.\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(text)}\n\n"
                f"CONTRACT:\n{compact_text(text, 60000)}"
            ),
        },
    ]
    try:
        return str(chat_json(client, provider, model, messages, retries=2).get("memory", "")).strip()
    except Exception as exc:
        log(f"术语记忆提取失败，将继续翻译。原因：{exc}")
        return ""


def build_chunks(items: list[dict], max_chars: int = 30000, max_paragraphs: int = 35) -> list[list[dict]]:
    chunks = []
    current = []
    current_chars = 0
    for item in items:
        item_chars = len(item["text"]) + sum(len(span["source_text"]) for span in item["format_spans"])
        if current and (current_chars + item_chars > max_chars or len(current) >= max_paragraphs):
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(item)
        current_chars += item_chars
    if current:
        chunks.append(current)
    return chunks


def translate_chunk(client: OpenAI, provider: str, model: str, memory: str, chunk: list[dict], context_before: str, context_after: str) -> dict:
    chunk_text = "\n".join(item.get("text", "") for item in chunk)
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Translate these Chinese legal contract paragraphs into precise formal legal English using the full chunk context.\n"
                "Then map every source formatting span to the exact English substring that should receive the same formatting.\n\n"
                "Mandatory paragraph-by-paragraph workflow:\n"
                "A. Treat each paragraph id as an independent formatting unit after using the chunk for translation context.\n"
                "B. For each paragraph, inspect the supplied run-derived format_spans as a checklist of formatted source content.\n"
                "C. For each format span, record the source formatted content, locate the legally corresponding content in that paragraph's English translation, and return that exact target_text.\n"
                "D. Complete this checklist for the paragraph before moving to the next paragraph.\n"
                "E. Do not rely on memory or visual guessing: every mapping must be grounded in the source span and the paragraph translation.\n\n"
                "Return JSON exactly like:\n"
                "{\"paragraphs\":[{\"id\":1,\"translation\":\"...\",\"format_mappings\":[{\"span_id\":\"p1_s0\",\"target_text\":\"...\",\"confidence\":\"high|medium|low\",\"note\":\"...\"}]}]}\n\n"
                "Rules:\n"
                "1. Translate each paragraph as a whole, not run-by-run, so the English is natural and legally accurate.\n"
                "2. Do not leave Chinese characters in translations unless they are intentionally part of a name or exhibit label.\n"
                "3. Every input format span must appear once in format_mappings.\n"
                "4. target_text must be an exact contiguous substring of the English translation whenever possible.\n"
                "5. If the formatted Chinese concept becomes non-contiguous in English, choose the closest legally equivalent contiguous phrase and mark confidence low.\n"
                "6. Preserve numbering, brackets, placeholders, article references, and draft notes accurately.\n\n"
                f"CAPITAL MARKETS LEGAL RAG:\n{legal_rag_for_text(chunk_text)}\n\n"
                f"TRANSLATION MEMORY:\n{memory or '(none)'}\n\n"
                f"CONTEXT BEFORE:\n{compact_text(context_before, 3000)}\n\n"
                f"CONTEXT AFTER:\n{compact_text(context_after, 3000)}\n\n"
                f"PARAGRAPHS JSON:\n{json.dumps({'paragraphs': chunk}, ensure_ascii=False)}"
            ),
        },
    ]
    return chat_json(client, provider, model, messages)


def translate_single_item(client: OpenAI, provider: str, model: str, memory: str, item: dict, context_before: str, context_after: str) -> dict:
    return translate_chunk(client, provider, model, memory, [item], context_before, context_after)


def cjk_count(text: str) -> int:
    return len(re.findall(r"[\u3400-\u9fff]", text or ""))


def needs_translation_repair(source_text: str, translation: str) -> bool:
    translation = translation or ""
    if not translation.strip():
        return True
    count = cjk_count(translation)
    if not count:
        return False
    compact_len = max(1, len(re.sub(r"\s+", "", translation)))
    source_count = max(1, cjk_count(source_text))
    ascii_letters = len(re.findall(r"[A-Za-z]", translation))
    if count >= 4 and count / compact_len >= 0.5 and ascii_letters < 10:
        return True
    return (count >= 20 and count / compact_len >= 0.08) or count >= max(30, source_count // 4)


def repair_chunk_results(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    chunk: list[dict],
    results: dict[int, dict],
    context_before: str,
    context_after: str,
    log,
) -> set[int]:
    repaired_ids: set[int] = set()
    bad_items = []
    for item in chunk:
        item_id = int(item["id"])
        result = results.get(item_id)
        translation = "" if not result else str(result.get("translation", "")).strip()
        if needs_translation_repair(item["text"], translation):
            bad_items.append(item)
    if not bad_items:
        return repaired_ids
    ids = ", ".join(str(item["id"] + 1) for item in bad_items[:12])
    more = "..." if len(bad_items) > 12 else ""
    log(f"自检查发现 {len(bad_items)} 个段落漏翻或中文残留，自动逐段重翻：{ids}{more}")
    for item in bad_items:
        item_id = int(item["id"])
        try:
            data = translate_single_item(client, provider, model, memory, item, context_before, context_after)
        except Exception as exc:
            log(f"段落 {item_id + 1} 自修复重翻失败：{exc}")
            continue
        for paragraph_result in data.get("paragraphs", []) or []:
            try:
                paragraph_id = int(paragraph_result.get("id"))
            except (TypeError, ValueError):
                continue
            if paragraph_id != item_id:
                continue
            translation = normalize_translation_against_source(item["text"], str(paragraph_result.get("translation", "")).strip())
            paragraph_result["translation"] = translation
            if translation:
                results[item_id] = paragraph_result
                repaired_ids.add(item_id)
            break
    return repaired_ids


def repair_all_results(
    client: OpenAI,
    provider: str,
    model: str,
    memory: str,
    items: list[dict],
    results: dict[int, dict],
    texts: list[str],
    log,
    progress,
    max_rounds: int = 5,
) -> set[int]:
    repaired_ids: set[int] = set()
    total = len(items)
    for round_index in range(1, max_rounds + 1):
        bad_items = []
        for item in items:
            item_id = int(item["id"])
            result = results.get(item_id)
            translation = "" if not result else str(result.get("translation", "")).strip()
            if needs_translation_repair(item["text"], translation):
                bad_items.append(item)
        if not bad_items:
            log(f"{APP_VERSION} 全文自检查通过：没有发现明显漏翻或大段中文残留。")
            return repaired_ids

        ids = ", ".join(str(int(item["id"]) + 1) for item in bad_items[:20])
        more = "..." if len(bad_items) > 20 else ""
        log(f"{APP_VERSION} 全文自检查第 {round_index} 轮：发现 {len(bad_items)} 个段落需要修复：{ids}{more}")
        changed = False
        for index, item in enumerate(bad_items, start=1):
            item_id = int(item["id"])
            progress(total - len(bad_items) + index - 1, total, f"全文自修复第 {round_index} 轮 {index}/{len(bad_items)}")
            context_before = "\n".join(texts[max(0, item_id - 5) : item_id])
            context_after = "\n".join(texts[item_id + 1 : min(len(texts), item_id + 6)])
            try:
                data = translate_single_item(client, provider, model, memory, item, context_before, context_after)
            except Exception as exc:
                log(f"段落 {item_id + 1} 自修复重翻失败：{exc}")
                continue
            for paragraph_result in data.get("paragraphs", []) or []:
                try:
                    paragraph_id = int(paragraph_result.get("id"))
                except (TypeError, ValueError):
                    continue
                if paragraph_id != item_id:
                    continue
                translation = normalize_translation_against_source(item["text"], str(paragraph_result.get("translation", "")).strip())
                paragraph_result["translation"] = translation
                if translation and translation != str(results.get(item_id, {}).get("translation", "")).strip():
                    results[item_id] = paragraph_result
                    repaired_ids.add(item_id)
                    changed = True
                break
        if not changed:
            remaining = [
                str(int(item["id"]) + 1)
                for item in items
                if needs_translation_repair(
                    item["text"], str(results.get(int(item["id"]), {}).get("translation", "")).strip()
                )
            ]
            log(
                f"{APP_VERSION} 全文自修复本轮没有取得新结果；将继续导出，并在 checklist 标出剩余疑似问题段落："
                f"{', '.join(remaining[:20])}{'...' if len(remaining) > 20 else ''}"
            )
            break
    return repaired_ids


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def nearest_nonspace(text: str, index: int, direction: int) -> str:
    while 0 <= index < len(text):
        if not text[index].isspace():
            return text[index]
        index += direction
    return ""


def candidate_match_score(text: str, start: int, end: int, exact_case: bool) -> int:
    score = 20 if exact_case else 0
    before = nearest_nonspace(text, start - 1, -1)
    after = nearest_nonspace(text, end, 1)
    if before in {'"', "'", "\u201c", "\u2018"} and after in {'"', "'", "\u201d", "\u2019"}:
        score += 100
    if before in {"(", "["} and after in {")", "]"}:
        score += 20
    if start == 0 or not text[start - 1].isalnum():
        score += 8
    else:
        score -= 50
    if end == len(text) or not text[end : end + 1].isalnum():
        score += 8
    else:
        score -= 50
    prefix = text[max(0, start - 8) : start].lower()
    if '"' in prefix or "\u201c" in prefix or "(" in prefix:
        score += 5
    return score


def find_available_span(text: str, needle: str, intervals: list[tuple[int, int]]) -> tuple[int, int] | None:
    if not needle:
        return None
    candidates = []
    for candidate in [needle, needle.strip(), normalize_spaces(needle), normalize_legal_english(needle)]:
        if candidate and candidate not in candidates:
            candidates.append(candidate)
    for candidate in candidates:
        if not candidate:
            continue
        matches = []
        start = text.find(candidate)
        while start != -1:
            end = start + len(candidate)
            if all(end <= a or start >= b for a, b in intervals):
                matches.append((candidate_match_score(text, start, end, True), start, end))
            start = text.find(candidate, start + 1)
        lower_text = text.lower()
        lower_candidate = candidate.lower()
        start = lower_text.find(lower_candidate)
        while start != -1:
            end = start + len(candidate)
            if all(end <= a or start >= b for a, b in intervals):
                matches.append((candidate_match_score(text, start, end, False), start, end))
            start = lower_text.find(lower_candidate, start + 1)
        if matches:
            matches.sort(key=lambda item: (-item[0], item[1]))
            _, start, end = matches[0]
            return start, end
    return None


def clear_paragraph_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)


def paragraph_rpr(paragraph):
    ppr = paragraph._p.pPr
    return None if ppr is None else ppr.find(qn("w:rPr"))


def base_format_overrides(paragraph) -> dict[str, bool]:
    overrides = {"bold": False, "italic": False, "underline": False, "highlight": False}
    rpr = paragraph_rpr(paragraph)
    if xml_has_off_property(rpr, "b"):
        overrides["bold"] = True
    if xml_has_off_property(rpr, "i"):
        overrides["italic"] = True
    if xml_has_off_property(rpr, "u"):
        overrides["underline"] = True
    if xml_has_off_property(rpr, "highlight"):
        overrides["highlight"] = True

    for run in paragraph.runs:
        if not run.text:
            continue
        if run.bold is False:
            overrides["bold"] = True
        if run.italic is False:
            overrides["italic"] = True
        if run.underline is False:
            overrides["underline"] = True
        if xml_has_off_property(run._r.rPr, "highlight"):
            overrides["highlight"] = True
    return overrides


def set_highlight_none(run) -> None:
    rpr = run._r.get_or_add_rPr()
    for existing in list(rpr.findall(qn("w:highlight"))):
        rpr.remove(existing)
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "none")
    rpr.append(highlight)


def apply_base_format(run, base_format: dict[str, bool]) -> None:
    if base_format.get("bold"):
        run.bold = False
    if base_format.get("italic"):
        run.italic = False
    if base_format.get("underline"):
        run.underline = False
    if base_format.get("highlight"):
        set_highlight_none(run)


def add_run_with_base_format(paragraph, text: str, base_format: dict[str, bool]):
    run = paragraph.add_run(text)
    apply_base_format(run, base_format)
    return run


def apply_span_format(run, span: FormatSpan, base_format: dict[str, bool]) -> None:
    apply_base_format(run, base_format)
    if span.bold:
        run.bold = True
    if span.italic:
        run.italic = True
    if span.underline:
        run.underline = span.underline
    if span.highlight_color:
        run.font.highlight_color = span.highlight_color


def rewrite_paragraph_with_format(paragraph, translation: str, intervals: list[dict]) -> None:
    intervals = sorted(intervals, key=lambda item: (item["start"], item["end"]))
    base_format = base_format_overrides(paragraph)
    clear_paragraph_runs(paragraph)
    cursor = 0
    for item in intervals:
        start = max(0, item["start"])
        end = min(len(translation), item["end"])
        if start < cursor or end <= start:
            continue
        if cursor < start:
            add_run_with_base_format(paragraph, translation[cursor:start], base_format)
        run = add_run_with_base_format(paragraph, translation[start:end], base_format)
        apply_span_format(run, item["span"], base_format)
        cursor = end
    if cursor < len(translation):
        add_run_with_base_format(paragraph, translation[cursor:], base_format)
    if not translation:
        add_run_with_base_format(paragraph, "", base_format)


def safe_save_document(doc: Document, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f"{output_path.stem}.tmp.{int(time.time() * 1000)}{output_path.suffix}")
    doc.save(str(tmp_path))
    try:
        os.replace(str(tmp_path), str(output_path))
        return output_path
    except OSError as exc:
        if getattr(exc, "winerror", None) != 5 and getattr(exc, "errno", None) != 13:
            tmp_path.unlink(missing_ok=True)
            raise
        fallback = output_path.with_name(f"{output_path.stem}_自动另存_{time.strftime('%Y%m%d_%H%M%S')}{output_path.suffix}")
        os.replace(str(tmp_path), str(fallback))
        return fallback


def build_markdown_checklist(rows: list[dict]) -> str:
    lines = [
        "# 格式映射 Checklist",
        "",
        "| 状态 | 段落 | 原格式文本 | 格式 | 英文映射 | 置信度 | 备注 |",
        "|---|---:|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append(
            "| {status} | {paragraph_id} | {source_text} | {style} | {target_text} | {confidence} | {note} |".format(
                status=row.get("status", ""),
                paragraph_id=row.get("paragraph_id", ""),
                source_text=str(row.get("source_text", "")).replace("|", "\\|"),
                style=str(row.get("style", "")).replace("|", "\\|"),
                target_text=str(row.get("target_text", "")).replace("|", "\\|"),
                confidence=str(row.get("confidence", "")).replace("|", "\\|"),
                note=str(row.get("note", "")).replace("|", "\\|"),
            )
        )
    return "\n".join(lines) + "\n"


def translate_docx_new_scheme(
    input_path: Path,
    output_dir: Path,
    provider: str,
    api_key: str,
    base_url: str,
    model: str,
    log,
    progress,
    cancel_event: threading.Event | None = None,
) -> tuple[Path, Path, Path, bool]:
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(str(input_path))
    paragraphs = list(iter_document_paragraphs(doc))
    texts = [paragraph_text(paragraph) for paragraph in paragraphs]
    targets = [idx for idx, text in enumerate(texts) if has_cjk(text)]
    if not targets:
        raise ValueError("没有在文档中发现中文内容。")

    client = build_client(api_key, base_url)
    memory = make_translation_memory(client, provider, model, texts, log)
    spans_by_id: dict[str, FormatSpan] = {}
    items = []
    for idx in targets:
        spans = extract_format_spans(paragraphs[idx], idx)
        public_spans = []
        for span in spans:
            spans_by_id[span.span_id] = span
            public_spans.append({"span_id": span.span_id, "source_text": span.source_text, "style": span.style_label})
        items.append({"id": idx, "text": texts[idx], "format_spans": public_spans})

    chunks = build_chunks(items)
    log(f"{APP_VERSION} 发现 {len(targets)} 个含中文段落，分为 {len(chunks)} 个大段翻译批次。")

    all_results = {}
    auto_repaired_paragraphs: set[int] = set()
    completed = 0
    cancelled = False
    for chunk_number, chunk in enumerate(chunks, start=1):
        if cancel_event and cancel_event.is_set():
            cancelled = True
            log("收到中止请求，停止进入下一批次，开始导出当前进度。")
            break
        first_id = chunk[0]["id"]
        last_id = chunk[-1]["id"]
        context_before = "\n".join(texts[max(0, first_id - 5) : first_id])
        context_after = "\n".join(texts[last_id + 1 : min(len(texts), last_id + 6)])
        progress(completed, len(targets), f"正在翻译第 {chunk_number}/{len(chunks)} 个大段...")
        log(f"[批次 {chunk_number}/{len(chunks)}] 段落 {first_id + 1} 到 {last_id + 1}，包含 {len(chunk)} 个待译段落。")
        try:
            data = translate_chunk(client, provider, model, memory, chunk, context_before, context_after)
        except Exception as exc:
            log(f"大段翻译失败，改为逐段处理该批次。原因：{exc}")
            data = {"paragraphs": []}
            for item in chunk:
                single = translate_single_item(client, provider, model, memory, item, context_before, context_after)
                data["paragraphs"].extend(single.get("paragraphs", []))
        for paragraph_result in data.get("paragraphs", []):
            all_results[int(paragraph_result.get("id"))] = paragraph_result
        completed += len(chunk)
        progress(completed, len(targets), f"第 {chunk_number}/{len(chunks)} 个大段已返回。")
        if cancel_event and cancel_event.is_set():
            cancelled = True
            log("当前批次已完成；根据中止请求，开始导出当前进度。")
            break

    if not cancelled:
        auto_repaired_paragraphs.update(
            repair_all_results(client, provider, model, memory, items, all_results, texts, log, progress)
        )

    checklist_rows = []
    if cancelled:
        checklist_rows.append(
            {
                "status": "CANCELLED",
                "paragraph_id": "",
                "source_text": "",
                "style": "",
                "target_text": "",
                "confidence": "",
                "note": "用户中止；已导出当前已完成批次，未完成段落保留原文。",
            }
        )
    raw_payload = {"input": items, "results": []}
    for item in items:
        paragraph_id = int(item["id"])
        result = all_results.get(paragraph_id)
        if not result:
            log(f"段落 {paragraph_id + 1} 缺少翻译结果，保留原文并写入 checklist。")
            for span in item["format_spans"]:
                checklist_rows.append(
                    {
                        "status": "MISSING_TRANSLATION",
                        "paragraph_id": paragraph_id + 1,
                        "source_text": span["source_text"],
                        "style": span["style"],
                        "target_text": "",
                        "confidence": "",
                        "note": "模型未返回该段翻译；若为中止导出，则该段尚未处理。",
                    }
                )
            continue

        translation = normalize_translation_against_source(item["text"], str(result.get("translation", "")).strip())
        if paragraph_id in auto_repaired_paragraphs:
            checklist_rows.append(
                {
                    "status": "AUTO_REPAIRED_TRANSLATION",
                    "paragraph_id": paragraph_id + 1,
                    "source_text": "",
                    "style": "",
                    "target_text": "",
                    "confidence": "deterministic",
                    "note": "自检查发现漏翻或中文残留后，已在导出前自动重翻该段。",
                }
            )
        mappings = result.get("format_mappings", []) or []
        mappings_by_span = {str(mapping.get("span_id")): mapping for mapping in mappings if mapping.get("span_id")}
        intervals = []
        occupied = []

        for public_span in item["format_spans"]:
            span_id = public_span["span_id"]
            source_span = spans_by_id[span_id]
            mapping = mappings_by_span.get(span_id)
            row = {
                "paragraph_id": paragraph_id + 1,
                "span_id": span_id,
                "source_text": public_span["source_text"],
                "style": public_span["style"],
                "target_text": "",
                "confidence": "",
                "note": "",
                "status": "",
            }
            if not mapping:
                fallback_found = None
                fallback_target = ""
                for candidate in source_span_fallback_targets(public_span["source_text"]):
                    fallback_target = normalize_legal_english(candidate)
                    fallback_found = find_available_span(translation, fallback_target, occupied)
                    if fallback_found:
                        break
                if fallback_found:
                    start, end = fallback_found
                    intervals.append({"start": start, "end": end, "span": source_span})
                    occupied.append((start, end))
                    row["target_text"] = fallback_target
                    row["confidence"] = "deterministic"
                    row["status"] = "APPLIED_FALLBACK"
                    row["note"] = "Applied deterministic legal-reference mapping because the model omitted this span."
                    checklist_rows.append(row)
                    continue
                row["status"] = "MISSING_MAPPING"
                row["note"] = "模型未返回该格式 span 的英文映射。"
                checklist_rows.append(row)
                continue
            target_text = normalize_legal_english(str(mapping.get("target_text", "")).strip())
            row["target_text"] = target_text
            row["confidence"] = str(mapping.get("confidence", "")).strip()
            row["note"] = str(mapping.get("note", "")).strip()
            found = find_available_span(translation, target_text, occupied)
            if not found:
                for candidate in source_span_fallback_targets(public_span["source_text"]):
                    target_text = normalize_legal_english(candidate)
                    found = find_available_span(translation, target_text, occupied)
                    if found:
                        row["target_text"] = target_text
                        row["confidence"] = "deterministic"
                        row["note"] = (row["note"] + " " if row["note"] else "") + "Used deterministic legal-reference fallback."
                        break
            if not found:
                row["status"] = "TARGET_NOT_FOUND"
                row["note"] = (row["note"] + " " if row["note"] else "") + "英文映射文本未在译文中精确找到。"
                checklist_rows.append(row)
                continue
            start, end = found
            intervals.append({"start": start, "end": end, "span": source_span})
            occupied.append((start, end))
            row["status"] = "APPLIED"
            checklist_rows.append(row)

        if needs_translation_repair(item["text"], translation):
            checklist_rows.append(
                {
                    "status": "HAS_CHINESE_IN_TRANSLATION",
                    "paragraph_id": paragraph_id + 1,
                    "source_text": "",
                    "style": "",
                    "target_text": "",
                    "confidence": "",
                    "note": "译文仍疑似存在漏翻或大段中文残留，请人工复核。",
                }
            )
        rewrite_paragraph_with_format(paragraphs[paragraph_id], translation, intervals)
        result["translation"] = translation
        raw_payload["results"].append(result)

    base_name = input_path.stem
    suffix = "_已中止" if cancelled else ""
    docx_path = output_dir / f"{base_name}_新方案英文翻译{suffix}.docx"
    checklist_path = output_dir / f"{base_name}_格式映射checklist{suffix}.md"
    json_path = output_dir / f"{base_name}_格式映射明细{suffix}.json"

    docx_path = safe_save_document(doc, docx_path)
    checklist_path.write_text(build_markdown_checklist(checklist_rows), encoding="utf-8")
    json_path.write_text(
        json.dumps({"cancelled": cancelled, "checklist": checklist_rows, "raw": raw_payload}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    progress(completed if cancelled else len(targets), len(targets), "已中止并导出当前进度" if cancelled else "全部完成")
    return docx_path, checklist_path, json_path, cancelled


class NewSchemeApp:
    def __init__(self, root):
        self.root = root
        self.main_thread_id = threading.get_ident()
        self.root.title("法律合同翻译新方案 - 大段翻译 + 格式映射")
        self.root.geometry("860x680")
        self.root.minsize(780, 620)

        config = load_config()
        provider = config.get("provider", "DeepSeek")
        if provider not in PROVIDER_DEFAULTS:
            provider = "DeepSeek"
        defaults = PROVIDER_DEFAULTS[provider]
        self.provider = StringVar(value=provider)
        self.api_key = StringVar(value=config.get("api_key", ""))
        self.base_url = StringVar(value=config.get("base_url", defaults["base_url"]))
        self.model = StringVar(value=config.get("model", defaults["model"]))
        self.file_path = StringVar(value="")
        self.output_dir = StringVar(value=str(OUTPUT_DIR))
        self.remember_key = BooleanVar(value=bool(config.get("api_key")))
        self.progress_value = StringVar(value="进度：0/0")
        self.progress_number = StringVar(value="0")
        self.current_text = StringVar(value="当前：尚未开始")
        self.cancel_event = threading.Event()

        top = Frame(root, padx=16, pady=12)
        top.pack(fill=X)
        Label(top, text="API 类型").pack(anchor="w")
        OptionMenu(top, self.provider, *PROVIDER_DEFAULTS.keys(), command=self.on_provider_change).pack(fill=X, pady=(4, 8))
        self.key_label = Label(top, text=defaults["key_label"])
        self.key_label.pack(anchor="w")
        Entry(top, textvariable=self.api_key, show="*", width=90).pack(fill=X, pady=(4, 8))
        Label(top, text="Base URL").pack(anchor="w")
        Entry(top, textvariable=self.base_url, width=90).pack(fill=X, pady=(4, 8))
        Label(top, text="模型").pack(anchor="w")
        Entry(top, textvariable=self.model, width=90).pack(fill=X, pady=(4, 8))
        Checkbutton(top, text="记住 API Key（明文保存在本机新方案目录）", variable=self.remember_key).pack(anchor="w")

        file_frame = Frame(root, padx=16, pady=8)
        file_frame.pack(fill=X)
        Label(file_frame, text="合同文件（.docx）").pack(anchor="w")
        row = Frame(file_frame)
        row.pack(fill=X, pady=(4, 8))
        Entry(row, textvariable=self.file_path).pack(side=LEFT, fill=X, expand=True)
        Button(row, text="选择文件", command=self.choose_file, width=12).pack(side=RIGHT, padx=(8, 0))
        Label(file_frame, text="输出目录").pack(anchor="w")
        out_row = Frame(file_frame)
        out_row.pack(fill=X, pady=(4, 0))
        Entry(out_row, textvariable=self.output_dir).pack(side=LEFT, fill=X, expand=True)
        Button(out_row, text="选择目录", command=self.choose_output_dir, width=12).pack(side=RIGHT, padx=(8, 0))

        self.drop_label = Label(
            root,
            text="把 Word 合同拖到这里\n输出：英文 DOCX + 格式映射 checklist + JSON 明细",
            relief="groove",
            height=4,
            padx=12,
            pady=12,
        )
        self.drop_label.pack(fill=X, padx=16, pady=8)
        if DND_AVAILABLE:
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind("<<Drop>>", self.on_drop)

        progress_frame = Frame(root, padx=16, pady=8)
        progress_frame.pack(fill=X)
        Label(progress_frame, textvariable=self.progress_value).pack(anchor="w")
        self.bar = Progressbar(progress_frame, maximum=100)
        self.bar.pack(fill=X, pady=(4, 6))
        Label(progress_frame, textvariable=self.current_text, wraplength=800, justify=LEFT).pack(anchor="w")

        action_frame = Frame(root, padx=16, pady=8)
        action_frame.pack(fill=X)
        self.start_button = Button(action_frame, text="开始新方案翻译", command=self.start, height=2)
        self.start_button.pack(side=LEFT, fill=X, expand=True)
        self.stop_button = Button(action_frame, text="中止并导出当前进度", command=self.request_cancel, height=2, state="disabled")
        self.stop_button.pack(side=RIGHT, padx=(8, 0))

        self.log_box = Text(root, height=14, wrap="word")
        self.log_box.pack(fill=BOTH, expand=True, padx=16, pady=(4, 16))
        self.log(f"{APP_VERSION} 新方案：大段上下文翻译 -> 大模型格式映射 -> checklist 防错漏。")
        self.log("注意：该方案会重建段落内 run，以便把原文粗体/下划线/高亮映射到英文对应短语。")

    def on_provider_change(self, provider):
        defaults = PROVIDER_DEFAULTS[provider]
        self.key_label.config(text=defaults["key_label"])
        self.base_url.set(defaults["base_url"])
        self.model.set(defaults["model"])

    def choose_file(self):
        path = filedialog.askopenfilename(title="选择 Word 合同", filetypes=[("Word 文档", "*.docx")])
        if path:
            self.file_path.set(path)

    def choose_output_dir(self):
        path = filedialog.askdirectory(title="选择输出目录", initialdir=self.output_dir.get() or str(OUTPUT_DIR))
        if path:
            self.output_dir.set(path)

    def on_drop(self, event):
        paths = self.root.tk.splitlist(event.data)
        if paths:
            self.file_path.set(paths[0])

    def log(self, text: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.log, text)
            return
        self.log_box.insert(END, text + "\n")
        self.log_box.see(END)
        self.root.update_idletasks()

    def update_progress(self, done: int, total: int, current: str):
        if threading.get_ident() != self.main_thread_id:
            self.root.after(0, self.update_progress, done, total, current)
            return
        percent = 0 if total <= 0 else max(0, min(100, done * 100 / total))
        self.bar["value"] = percent
        self.progress_value.set(f"进度：{done}/{total}（{percent:.1f}%）")
        self.current_text.set(f"当前：{current}")
        self.root.update_idletasks()

    def start(self):
        provider = self.provider.get().strip()
        api_key = self.api_key.get().strip()
        base_url = self.base_url.get().strip()
        model = self.model.get().strip()
        input_path = Path(self.file_path.get().strip().strip('"'))
        output_dir = Path(self.output_dir.get().strip().strip('"'))
        if not api_key:
            messagebox.showerror("缺少 API Key", f"请先输入 {provider} API Key。")
            return
        if not model:
            messagebox.showerror("缺少模型", "请填写模型名称。")
            return
        if not input_path.exists() or input_path.suffix.lower() != ".docx":
            messagebox.showerror("文件不支持", "请选择存在的 .docx 文件。")
            return
        if provider == "DeepSeek" and not base_url:
            base_url = PROVIDER_DEFAULTS["DeepSeek"]["base_url"]
            self.base_url.set(base_url)
        if self.remember_key.get():
            save_config(provider, api_key, base_url, model)
        self.cancel_event.clear()
        self.start_button.config(state="disabled")
        self.stop_button.config(state="normal")
        thread = threading.Thread(target=self.worker, args=(input_path, output_dir, provider, api_key, base_url, model), daemon=True)
        thread.start()

    def request_cancel(self):
        self.cancel_event.set()
        self.stop_button.config(state="disabled")
        self.log("已请求中止；当前 API 批次返回后会导出当前进度。")
        self.update_progress(0, 0, "正在等待当前批次返回，然后导出当前进度...")

    def worker(self, input_path: Path, output_dir: Path, provider: str, api_key: str, base_url: str, model: str):
        try:
            docx_path, checklist_path, json_path, cancelled = translate_docx_new_scheme(
                input_path, output_dir, provider, api_key, base_url, model, self.log, self.update_progress, self.cancel_event
            )
            self.root.after(0, self.finish_success, docx_path, checklist_path, json_path, cancelled)
        except Exception as exc:
            self.log("发生错误：")
            self.log(str(exc))
            self.log(traceback.format_exc())
            self.root.after(0, self.finish_error, str(exc))

    def finish_success(self, docx_path: Path, checklist_path: Path, json_path: Path, cancelled: bool):
        self.start_button.config(state="normal")
        self.stop_button.config(state="disabled")
        title = "已中止并导出" if cancelled else "完成"
        messagebox.showinfo(title, f"已输出：\n{docx_path}\n{checklist_path}\n{json_path}")

    def finish_error(self, error: str):
        self.start_button.config(state="normal")
        self.stop_button.config(state="disabled")
        messagebox.showerror("翻译失败", error)


def main():
    root_class = TkinterDnD.Tk if DND_AVAILABLE else Tk
    root = root_class()
    NewSchemeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
